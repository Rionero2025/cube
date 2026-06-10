import calendar
import hashlib
import hmac
import html
import json
import re
import sqlite3
from datetime import date, datetime, timedelta
from io import BytesIO
from pathlib import Path
from typing import Any

import pandas as pd
import streamlit as st
from docx import Document
from docx.shared import Pt, Inches
from reportlab.lib import colors
from reportlab.lib.enums import TA_CENTER, TA_LEFT
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import cm
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, PageBreak

BASE_DIR = Path(__file__).resolve().parent
DB_PATH = BASE_DIR / "data" / "cube_contracts_pro.db"
LOGO_PATH = BASE_DIR / "assets" / "logo_cube.png"
GENERATED_DIR = BASE_DIR / "generated_contracts"
DOC_DIR = BASE_DIR / "allegati" / "documenti"
WORK_DIR = BASE_DIR / "allegati" / "lavori"
PAY_DIR = BASE_DIR / "allegati" / "pagamenti"
INV_DIR = BASE_DIR / "allegati" / "fatture"
REPORT_DIR = BASE_DIR / "resoconti_lavori"

for p in [GENERATED_DIR, DOC_DIR, WORK_DIR, PAY_DIR, INV_DIR, REPORT_DIR, DB_PATH.parent]:
    p.mkdir(parents=True, exist_ok=True)

st.set_page_config(page_title="CUBE Management Contract", page_icon="📘", layout="wide", initial_sidebar_state="expanded")

# -----------------------------
# Utility
# -----------------------------
def today_iso() -> str:
    return date.today().isoformat()

def now_iso() -> str:
    return datetime.now().isoformat(timespec="seconds")

def add_months(d: date, months: int) -> date:
    month = d.month - 1 + months
    year = d.year + month // 12
    month = month % 12 + 1
    day = min(d.day, calendar.monthrange(year, month)[1])
    return date(year, month, day)

def parse_date(v: Any, default: date | None = None) -> date | None:
    if isinstance(v, date):
        return v
    if v in (None, "", pd.NaT):
        return default
    try:
        return datetime.strptime(str(v)[:10], "%Y-%m-%d").date()
    except Exception:
        return default


def to_date(v: Any, default: date | None = None) -> date | None:
    return parse_date(v, default)

def money(v: Any) -> str:
    try:
        return f"€ {float(v):,.2f}".replace(",", "X").replace(".", ",").replace("X", ".")
    except Exception:
        return "€ 0,00"

def safe(s: Any) -> str:
    return html.escape(str(s or ""))

def slug(s: str) -> str:
    s = re.sub(r"[^A-Za-z0-9_\-]+", "_", str(s or "file")).strip("_")
    return s[:80] or "file"



def hash_password(password: str) -> str:
    raw = ("CUBE_MANAGEMENT_CONTRACT|" + str(password or "")).encode("utf-8")
    return hashlib.sha256(raw).hexdigest()

def verify_password(password: str, password_hash: str | None) -> bool:
    if not password_hash:
        return False
    return hmac.compare_digest(hash_password(password), str(password_hash))

def current_staff_user() -> dict | None:
    staff_id = st.session_state.get("staff_user_id")
    if not staff_id:
        return None
    df = read_df("SELECT * FROM staff WHERE id=?", (staff_id,))
    if df.empty:
        st.session_state.pop("staff_user_id", None)
        return None
    return df.iloc[0].to_dict()

def status_badge_html(stato: str, residuo: float = 0.0, data_scadenza: str | None = None) -> str:
    stato = str(stato or "Da pagare")
    try:
        is_scaduta = bool(data_scadenza and str(data_scadenza) < date.today().isoformat() and float(residuo or 0) > 0.01 and stato != "Pagata")
    except Exception:
        is_scaduta = False

    if is_scaduta:
        label, bg, fg, border = "Scaduta", "#ffe8e8", "#b42318", "#f5b5b5"
    elif stato == "Pagata":
        label, bg, fg, border = "Pagata", "#e8f8ee", "#087443", "#b7e4c7"
    elif stato == "Acconto":
        label, bg, fg, border = "Acconto", "#fff2df", "#b45309", "#ffd08a"
    elif stato == "Sollecitata":
        label, bg, fg, border = "Sollecitata", "#fff4cc", "#92400e", "#facc15"
    elif stato == "Annullata":
        label, bg, fg, border = "Annullata", "#eeeeee", "#525252", "#d4d4d4"
    else:
        label, bg, fg, border = "Da pagare", "#eaf3ff", "#0b5cab", "#b9d8ff"

    return f"<span style='display:inline-block;padding:5px 10px;border-radius:999px;background:{bg};color:{fg};border:1px solid {border};font-weight:700;font-size:12px'>{label}</span>"


def render_payment_simple_table(df: pd.DataFrame, key_prefix: str) -> int | None:
    if df.empty:
        st.info("Nessuna rata da mostrare.")
        return None

    visual = df.copy()
    visual["Cliente"] = visual["cliente"].astype(str) if "cliente" in visual.columns else ""
    visual["Rata"] = visual["numero_rata"].astype(int)
    visual["Scadenza"] = visual["data_scadenza"].astype(str)
    visual["Importo rata"] = visual["totale"].apply(money)
    visual["Pagato"] = visual["pagato"].apply(money)
    visual["Residuo"] = visual["residuo"].apply(money)
    visual["Stato"] = visual.apply(lambda r: status_badge_html(r.get("stato"), r.get("residuo"), r.get("data_scadenza")), axis=1)

    cols = ["Cliente","Rata","Scadenza","Importo rata","Pagato","Residuo","Stato"] if visual["Cliente"].nunique() > 1 else ["Rata","Scadenza","Importo rata","Pagato","Residuo","Stato"]
    st.markdown(
        visual[cols].to_html(escape=False, index=False),
        unsafe_allow_html=True
    )

    opts = {
        f"Rata {int(r['numero_rata'])} · {r['data_scadenza']} · residuo {money(r['residuo'])} · ID {int(r['id'])}": int(r["id"])
        for _, r in df.iterrows()
    }
    return opts[st.selectbox("Seleziona rata da gestire", list(opts.keys()), key=f"simple_rata_{key_prefix}")]

def save_upload(uploaded, folder: Path, prefix: str = "") -> str | None:
    if not uploaded:
        return None
    folder.mkdir(parents=True, exist_ok=True)
    fname = f"{datetime.now().strftime('%Y%m%d_%H%M%S')}_{prefix}{slug(uploaded.name)}"
    path = folder / fname
    path.write_bytes(uploaded.getvalue())
    return str(path.relative_to(BASE_DIR))

def file_download_button(rel_path: str | None, label: str = "Scarica allegato"):
    if not rel_path:
        return
    path = BASE_DIR / rel_path
    if path.exists():
        st.download_button(label, data=path.read_bytes(), file_name=path.name, key=f"dl_{rel_path}_{datetime.now().microsecond}")


def get_upload_bytes(uploaded: Any) -> bytes:
    if uploaded is None:
        return b""
    if isinstance(uploaded, bytes):
        return uploaded
    if hasattr(uploaded, "getvalue"):
        return uploaded.getvalue()
    if hasattr(uploaded, "read"):
        data = uploaded.read()
        try:
            uploaded.seek(0)
        except Exception:
            pass
        return data
    return b""

def _repair_spaced_letters(text: str) -> str:
    """Ricompone parole estratte come C O N C E T T I o S a r t o r i a l i."""
    def repl(m):
        s = m.group(0)
        joined = re.sub(r"\s+", "", s)
        # Evita di fondere sequenze numeriche/date; utile soprattutto per nomi e titoli.
        return joined if len(joined) >= 4 else s
    # solo sequenze di lettere singole separate da spazi
    return re.sub(r"(?<![A-Za-zÀ-Üà-ü])(?:[A-Za-zÀ-Üà-ü]\s+){3,}[A-Za-zÀ-Üà-ü](?![A-Za-zÀ-Üà-ü])", repl, text)

def _add_spaces_around_labels(text: str) -> str:
    labels = [
        "Ragione Sociale", "Indirizzo Sede Legale", "Sede Legale", "Città", "Citta", "CAP",
        "Provincia", "Stato", "P.IVA", "PIVA", "Cod. Fisc", "Codice Fiscale", "Referente",
        "Telefono ufficio", "Cellulare", "E-mail", "Email", "Data stipula contratto", "Protocollo",
        "Servizi offerti", "Totale Servizi Acquistati", "Modalità di pagamento", "Modalita di pagamento",
        "Luogo e data", "Corrispettivo", "Compenso", "Foro", "PEC", "REA", "Codice Ateco"
    ]
    for lab in labels:
        compact = re.sub(r"\s+", "", lab)
        text = re.sub(compact, lab, text, flags=re.I)
    return text

def compact_pdf_text_repair(text: str) -> str:
    if not text:
        return ""
    text = str(text).replace("\x00", " ").replace("\r", "\n")
    text = text.replace("□", " ").replace("☒", " [x] ").replace("", " • ")
    text = text.replace("—", " - ").replace("–", " - ")
    # Molti PDF firmati/convertiti producono parole come C O N C E T T I o addirittura intere righe
    # con una lettera alla volta. Prima ricomponiamo quelle sequenze, poi normalizziamo le etichette.
    text = _repair_spaced_letters(text)
    # Fallback più aggressivo: se una riga contiene tante lettere singole, le riuniamo in blocchi leggibili.
    fixed_lines = []
    for line in text.splitlines():
        tokens = line.split()
        if len(tokens) >= 8 and sum(1 for t in tokens if re.fullmatch(r"[A-Za-zÀ-Üà-ü]", t)) / max(1, len(tokens)) > 0.55:
            line = re.sub(r"(?<!\w)((?:[A-Za-zÀ-Üà-ü]\s+){3,}[A-Za-zÀ-Üà-ü])(?!\w)", lambda m: m.group(1).replace(" ", ""), line)
        fixed_lines.append(line)
    text = "\n".join(fixed_lines)
    # ripara parole tipiche estratte senza spazio dai moduli grafici
    fixes = {
        "RagioneSociale": "Ragione Sociale", "SedeLegale": "Sede Legale", "IndirizzoSedeLegale": "Indirizzo Sede Legale",
        "Cod.Fisc.": "Cod. Fisc.", "CodFisc": "Cod. Fisc.", "P.IVA/C.F.": "P.IVA/C.F.", "P.IVA/Cod.Fisc.": "P.IVA/Cod. Fisc.",
        "Telefonoufficio": "Telefono ufficio", "Datastipulacontratto": "Data stipula contratto",
        "Svilupposito": "Sviluppo sito", "Aperturamercato": "Apertura mercato", "Gestionecanale": "Gestione canale",
        "Consulenzaeanalisi": "Consulenza e analisi", "GestioneADS": "Gestione ADS", "Riunione mensile": "Riunione mensile",
        "Gestioneinserimento": "Gestione inserimento", "UsoIA": "Uso IA", "Creazionecontenuti": "Creazione contenuti",
        "Generazioneimmagini": "Generazione immagini", "InserimentoVetrine": "Inserimento vetrine", "Marketplace": "Marketplace",
    }
    for a, b in fixes.items():
        text = re.sub(re.escape(a), b, text, flags=re.I)
    text = _add_spaces_around_labels(text)
    text = re.sub(r"(?<=[a-zàèéìòù])(?=[A-ZÀ-Ü])", " ", text)
    text = re.sub(r"[ \t]+", " ", text)
    text = re.sub(r"\n{3,}", "\n\n", text)
    return text.strip()

def clean_import_preview_text(raw: str) -> str:
    txt = compact_pdf_text_repair(raw)
    txt = re.sub(r"\s*([.;:])\s*", r"\1 ", txt)
    txt = re.sub(r"\s+", " ", txt).strip()
    return txt

def _ocr_image_bytes(data: bytes) -> str:
    try:
        from PIL import Image, ImageOps, ImageEnhance, ImageFilter
        import pytesseract
        img = Image.open(BytesIO(data)).convert("RGB")
        # migliora foto storte/scure: scala, grigio, contrasto e threshold leggero
        max_w = 2200
        if img.width < max_w:
            ratio = max_w / max(1, img.width)
            img = img.resize((int(img.width * ratio), int(img.height * ratio)))
        gray = ImageOps.grayscale(img)
        gray = ImageEnhance.Contrast(gray).enhance(1.8)
        gray = gray.filter(ImageFilter.SHARPEN)
        # config pensato per moduli/contratti con righe e testo misto
        cfg = "--oem 3 --psm 6"
        txt1 = pytesseract.image_to_string(gray, lang="ita+eng", config=cfg)
        # secondo passaggio più libero per documenti interi
        txt2 = pytesseract.image_to_string(gray, lang="ita+eng", config="--oem 3 --psm 11")
        return compact_pdf_text_repair(txt1 + "\n" + txt2)
    except Exception as exc:
        raise RuntimeError("Per leggere foto/scansioni serve Tesseract OCR installato sul PC oltre a pytesseract. Il sistema resta comunque editabile manualmente.") from exc

def _pdf_text_with_pypdf(data: bytes) -> str:
    from pypdf import PdfReader
    reader = PdfReader(BytesIO(data))
    chunks = []
    for page in reader.pages:
        try:
            chunks.append(page.extract_text() or "")
        except TypeError:
            chunks.append(page.extract_text() or "")
        except Exception:
            pass
        # alcune versioni supportano layout mode
        try:
            chunks.append(page.extract_text(extraction_mode="layout") or "")
        except Exception:
            pass
    return "\n".join(chunks)

def _pdf_text_with_pymupdf(data: bytes) -> str:
    try:
        import fitz  # PyMuPDF
    except Exception:
        return ""
    chunks = []
    doc = fitz.open(stream=data, filetype="pdf")
    for page in doc:
        try:
            chunks.append(page.get_text("text", sort=True) or "")
        except Exception:
            pass
        try:
            blocks = page.get_text("blocks", sort=True) or []
            blocks = sorted(blocks, key=lambda b: (round(b[1] / 8), b[0]))
            chunks.append("\n".join(str(b[4]) for b in blocks if len(b) > 4))
        except Exception:
            pass
    return "\n".join(chunks)

def _pdf_ocr_with_pymupdf(data: bytes, max_pages: int = 3) -> str:
    try:
        import fitz
        from PIL import Image
    except Exception:
        return ""
    chunks = []
    doc = fitz.open(stream=data, filetype="pdf")
    for i, page in enumerate(doc):
        if i >= max_pages:
            break
        try:
            pix = page.get_pixmap(matrix=fitz.Matrix(2.5, 2.5), alpha=False)
            img = Image.frombytes("RGB", [pix.width, pix.height], pix.samples)
            bio = BytesIO()
            img.save(bio, format="PNG")
            chunks.append(_ocr_image_bytes(bio.getvalue()))
        except Exception:
            pass
    return "\n".join(chunks)

def _quality_score_text(t: str) -> int:
    if not t:
        return 0
    score = len(re.findall(r"[A-Za-zÀ-Üà-ü]{3,}", t))
    score += 5 * len(re.findall(r"\b\d{11}\b", t))
    score += 8 * len(re.findall(r"[A-Za-z0-9._%+\-]+@[A-Za-z0-9.\-]+\.[A-Za-z]{2,}", t))
    score += 10 if re.search(r"corrispettivo|compenso|servizi|ragione sociale|contratto", t, re.I) else 0
    return score

def estrai_testo_da_upload_contratto(uploaded: Any) -> str:
    name = (getattr(uploaded, "name", "") or "").lower()
    data = get_upload_bytes(uploaded)
    if not data:
        return ""
    texts = []
    if name.endswith(".pdf"):
        # 1) lettura testuale classica
        try:
            texts.append(_pdf_text_with_pypdf(data))
        except Exception:
            pass
        # 2) lettura per blocchi/coordinate, spesso migliore nei PDF grafici
        try:
            texts.append(_pdf_text_with_pymupdf(data))
        except Exception:
            pass
        combined = "\n".join(t for t in texts if t)
        # 3) OCR solo se il testo è scarso o troppo incollato/spaziato
        score = _quality_score_text(combined)
        spaced_noise = bool(re.search(r"(?:[A-Za-z]\s+){8,}", combined or ""))
        # OCR è costoso: lo usiamo solo quando la lettura testuale è davvero insufficiente.
        if score < 25 or (score < 80 and spaced_noise):
            try:
                texts.append(_pdf_ocr_with_pymupdf(data, max_pages=2))
            except Exception:
                pass
        final = compact_pdf_text_repair("\n".join(t for t in texts if t))
        if not final:
            raise RuntimeError("Non sono riuscito a leggere il PDF. Prova a caricare una scansione più nitida o una foto frontale.")
        return final
    # Foto / scansioni
    return _ocr_image_bytes(data)

def title_case_societa(value: str) -> str:
    value = re.sub(r"\s+", " ", str(value or "")).strip(" .,:;-")
    if not value:
        return ""
    # separa alcune parole quando il PDF le incolla tra loro
    value = re.sub(r"(?<=[a-zàèéìòù])(?=[A-ZÀ-Ü])", " ", value)
    value = value.replace("Società", "SOCIETÀ").replace("Responsabilità", "RESPONSABILITÀ").replace("Limitata", "LIMITATA")
    # normalizza le sigle societarie più comuni
    repl = {
        "Srl": "S.R.L.", "S.r.l.": "S.R.L.", "Srls": "S.R.L.S.", "S.r.l.s.": "S.R.L.S.",
        "Sas": "S.A.S.", "Spa": "S.P.A.", "Semplificata": "SEMPLIFICATA"
    }
    parts = []
    for w in value.split():
        clean = w.strip()
        parts.append(repl.get(clean, clean.capitalize() if clean.islower() else clean))
    return " ".join(parts).strip()

def _first(pattern: str, text: str, default: str = "", flags=re.I|re.S) -> str:
    m = re.search(pattern, text or "", flags)
    return (m.group(1).strip(" .,:;\n\t") if m else default)

def _parse_euro_num(v: str) -> float:
    if not v: return 0.0
    s = str(v).replace("€", "").replace(" ", "")
    if "," in s: s = s.replace(".", "").replace(",", ".")
    try: return float(s)
    except Exception: return 0.0

def _parse_it_date(v: str) -> date | None:
    if not v: return None
    v = v.strip().replace("-", "/")
    for fmt in ["%d/%m/%Y", "%d/%m/%y", "%Y/%m/%d"]:
        try: return datetime.strptime(v, fmt).date()
        except Exception: pass
    return None

def estrai_dati_contratto_da_testo(raw_text: str) -> dict:
    text = compact_pdf_text_repair(raw_text)
    emails_raw = re.findall(r"[A-Za-z0-9._%+\-]+@[A-Za-z0-9.\-]+\.[A-Za-z]{2,}", text)
    emails = []
    for e in emails_raw:
        # nei moduli OCR il numero cellulare può restare attaccato davanti all'email
        e = re.sub(r"^\d{6,}", "", e).strip(" .,:;")
        emails.append(e)
    email_cliente = next((e for e in emails if "bebol" not in e.lower() and "cube" not in e.lower()), "")
    pive = re.findall(r"\b\d{11}\b", text)
    piva_cliente = next((x for x in pive if x not in {"09814141215"}), "")
    tel = _first(r"(?:Telefono\s*ufficio|Tel\.?|Cellulare|Cel\.?)\s*[:\-]?\s*(\+?\d[\d\s]{6,18})", text)
    if not tel:
        nums = [n for n in re.findall(r"\b3\d{8,10}\b", text) if n != piva_cliente]
        tel = nums[0] if nums else ""

    # Modulo grafico BEBOL: Ragione sociale, indirizzo, città, PIVA, servizi e rate.
    is_modulo_bebol = bool(re.search(r"Contratto\s+tra\s+le\s+parti|Protocollo|Totale\s+Servizi\s+Acquistati|SEPA\s+Direct|Totale\s+pacchetto\s+SMS", text, re.I))
    # I contratti lunghi con articoli 1-15 non sono moduli grafici anche se contengono tabelle/servizi.
    if re.search(r"ART\.?\s*1|Art\.?\s*1", text) and re.search(r"CONTRATTO\s+DI\s+CONSULENZA", text, re.I):
        is_modulo_bebol = False
    if (not is_modulo_bebol) and re.search(r"Sviluppo\s+sito|Apertura\s+mercato|Gestione\s+inserimento|Gestione\s+social", text, re.I) and re.search(r"(?:[3-9]00|[3-9]00[3-9]00)", re.sub(r"\s+", "", text)):
        is_modulo_bebol = True
    rag = ""; sede = ""; cf = ""; data_firma = None; importo = 0.0; servizi = []
    def _extract_rate_values_from_text(t: str) -> list[int]:
        vals = []
        # importi espliciti vicino a Rata/Acconto
        for m in re.finditer(r"(?:Rata|Acconto)\s*€?\s*([0-9]{3,4})", t, re.I):
            try: vals.append(int(m.group(1)))
            except Exception: pass
        # sequenze di tre cifre anche incollate: 400400400500 -> 400,400,400,500
        for block in re.findall(r"(?<!\d)(?:[3-9]00){2,}(?!\d)", re.sub(r"\s+", "", t)):
            for i in range(0, len(block), 3):
                try: vals.append(int(block[i:i+3]))
                except Exception: pass
        # importi isolati plausibili
        for n in re.findall(r"\b([3-9]00|1[0-9]00|2[0-9]00|3[0-9]00|4[0-9]00|5[0-9]00)\b", t):
            try: vals.append(int(n))
            except Exception: pass
        # elimina duplicati e falsi positivi, ma conserva l'ordine
        out = []
        for v in vals:
            if 100 <= v <= 5000:
                out.append(v)
        return out

    if is_modulo_bebol:
        rag = _first(r"Ragione\s*Sociale\s*[:\-]?\s*(.+?)\s+Indirizzo", text)
        if not rag:
            rag = _first(r"^\s*([A-ZÀ-Üa-zà-ü0-9'’\. ]{3,120})\s+Via\s", text)
        if not rag:
            rag = _first(r"([A-ZÀ-Üa-zà-ü'’ ]{3,120})\s+Via\s+[A-ZÀ-Üa-zà-ü]", text)
        rag = re.sub(r"^\d{6,}", "", rag or "").replace("Ragione Sociale", "").strip(" .,:;-")
        rag = title_case_societa(rag)
        via = _first(r"Indirizzo\s*Sede\s*Legale\s*[:\-]?\s*(.+?)\s+Citt", text)
        if not via: via = _first(r"(Via\s+.+?)\s+(?:Campobasso|Napoli|Roma|Milano|Firenze|Caserta|Salerno|Siano)", text)
        if not via and rag:
            via = _first(rf"{re.escape(rag).replace('\\ ', r'\s+')}\s+(Via\s+.+?)(?:\s+\d{{11}}|\s+Referente|\s+Telefono|\s+Cellulare|$)", text, flags=re.I|re.S)
        citta = _first(r"Citt[àa]\s*[:\-]?\s*(.+?)\s*[-–]?\s*CAP", text)
        cap = _first(r"CAP\s*[:\-]?\s*(\d{5})", text)
        prov = _first(r"Provincia\s*[:\-]?\s*([A-Z]{2})", text)
        stato = _first(r"Stato\s*[:\-]?\s*([A-ZÀ-Üa-zà-ü ]+?)\s*-?\s*P\.?IVA", text) or "Italia"
        if (not citta or not cap or not prov) and via:
            # Esempi: "Via Marconi n. 16Campobasso86100 CB Italia" oppure "Via Marconi n. 16 Campobasso 86100 CB Italia"
            maddr = re.search(r"(Via\s+[A-ZÀ-Üa-zà-ü0-9\.\,\/\- ]{1,70}?)([A-ZÀ-Ü][a-zà-ü'’ ]{3,40})\s*(\d{5})\s+([A-Z]{2})\s+(Italia)", text, re.I|re.S)
            if maddr:
                via, citta, cap, prov, stato = [maddr.group(i).strip() for i in range(1,6)]
        sede = ", ".join([x for x in [via, citta, cap, prov, stato] if x])
        if piva_cliente: cf = piva_cliente
        # data dal modulo: Data stipula contratto o sequenza giorno mese anno
        data_firma = _parse_it_date(_first(r"Data\s+stipula\s+contratto\s*[:\-]?\s*(\d{1,2}[/-]\d{1,2}[/-]\d{2,4})", text))
        if data_firma is None:
            # es. Protocollo GB 06 26 - Data 28 04 2026
            m = re.search(r"\b(\d{1,2})\s+(\d{1,2})\s+(20\d{2})\b", text)
            if m:
                data_firma = _parse_it_date(f"{m.group(1)}/{m.group(2)}/{m.group(3)}")
        # servizi: parte testuale prima delle rate, sia moduli compilati a PC sia foto con scrittura/OCR
        blocco = _first(r"Servizi\s+offerti.*?(.+?)(?:Totale\s+Servizi|Totale\s+pacchetto|Modalit[àa]|Pagamento\s+mensile|SEPA|\b[3-9]00\b)", text)
        if not blocco:
            blocco = _first(r"(Sviluppo\s+sito.+?Riunione\s+mensile)", text)
        if blocco:
            keys = r"Sviluppo\s+sito(?:\s+web)?|Apertura\s+(?:e-?commerce|mercato)|Gestione\s+social|Gestione\s+canale|Gestione\s+inserimento|Creazione\s+contenuti|Generazione\s+immagini|Inserimento\s+(?:vetrine|vestiti)|Uso\s+IA|Consulenza\s+e\s+analisi|Gestione\s+ADS|Riunione\s+mensile|Marketplace|Dopo\s+\d+\s+mesi"
            matches = list(re.finditer(keys, blocco, re.I))
            if matches:
                for idx, m in enumerate(matches):
                    part_start = m.start()
                    part_end = matches[idx+1].start() if idx + 1 < len(matches) else len(blocco)
                    part = blocco[part_start:part_end]
                    part = re.sub(r"\s+", " ", part).strip(" .;,-")
                    if len(part) > 2:
                        servizi.append({"Servizio": part[:90], "Descrizione": part})
            else:
                for part in re.split(r"[\n;•]+", blocco):
                    part = re.sub(r"\s+", " ", part).strip(" .;,-")
                    if len(part) > 3:
                        servizi.append({"Servizio": part[:90], "Descrizione": part})
        rate_vals = _extract_rate_values_from_text(text)
        if rate_vals:
            # Nei PDF/foto l'OCR può duplicare le rate: l'ultima sequenza completa è di solito la più pulita.
            seq = rate_vals[-12:] if len(rate_vals) >= 12 else rate_vals
            importo = float(sum(seq)) if len(seq) >= 2 else float(seq[0])
    else:
        # Contratto testuale standard: cerca prima la parte dopo "e"/"con" e prima della P.IVA del cliente.
        rag = _first(r"Cliente:\s*([A-Z0-9À-Ü'’&\.\- ]{5,160})(?:\s+PARTI|\s+TRA|$)", text)
        if not rag:
            rag = _first(r"\be\s+([^,]{5,180}?)(?:,\s*P\.?IVA|,\s*VAT|,\s*con sede|,\s*società)", text, flags=re.I|re.S)
        if not rag:
            rag = _first(r"\bcon\s+([^,]{5,180}?)(?:,\s*P\.?IVA|,\s*VAT)", text, flags=re.I|re.S)
        rag = title_case_societa(rag)

        # Sede cliente: prende la sede che segue il nome cliente, fino a REA/PEC/titolare/virgola forte.
        if rag:
            rag_pattern = re.escape(rag).replace("\\ ", r"\s+")
            sede = _first(rf"{rag_pattern}.*?con sede legale in\s+(.+?)(?:,\s*REA|,\s*PEC|,\s*titolare|;|\. )", text, flags=re.I|re.S)
        sedi = re.findall(r"con sede(?:\s+legale)?\s+in\s+(.+?)(?:,\s*P\.?\s*IVA|,\s*IVA|,\s*REA|,\s*PEC|,\s*titolare|,\s*in persona|;|\. )", text, re.I|re.S)
        sedi_clean = [re.sub(r"\s+", " ", s).strip(" .,:;") for s in sedi]
        cliente_sedi = [s for s in sedi_clean if "Benedetto Brin" not in s and "80142 Napoli" not in s]
        if cliente_sedi:
            sede = cliente_sedi[-1]
        elif not sede and sedi_clean:
            sede = sedi_clean[-1]

        cf_matches = []
        for mcf in re.finditer(r"C\s*\.?\s*F\s*\.?\s*[:\-]?\s*([A-Z0-9\s]{11,28})", text, re.I):
            val = re.sub(r"\s+", "", mcf.group(1)).strip(" .,:;-")
            if val and val not in {"09814141215"} and val != piva_cliente:
                cf_matches.append(val)
        cf = cf_matches[0] if cf_matches else _first(r"(?:C\.?F\.?|CF)\s*([A-Z0-9]{11,16})", text)
        if cf == "09814141215":
            cf = piva_cliente
        data_firma = _parse_it_date(_first(r"Luogo\s+e\s+data:\s*[^,]*,\s*(\d{1,2}[/-]\d{1,2}[/-]\d{4})", text)) or _parse_it_date(_first(r"(\d{1,2}[/-]\d{1,2}[/-]\d{4})", text))
        # Corrispettivo / compenso: preferisce sempre il totale annuo/complessivo.
        # Esempio MUNDI: "compenso di € 400,00 + IVA al mese, per 12 mesi, pari a € 4.800,00 + IVA annui".
        importo = _parse_euro_num(_first(r"pari\s+a\s*€\s*([0-9\.\,]+)\s*(?:\([^)]*\))?\s*\+?\s*IVA?\s*(?:annui|annuo|annuale)", text, flags=re.I|re.S))
        if not importo:
            importo = _parse_euro_num(_first(r"(?:corrispettivo|compenso)\s+complessivo\s+(?:pari\s+a|di)?\s*€\s*([0-9\.\,]+)", text, flags=re.I|re.S))
        if not importo:
            importo = _parse_euro_num(_first(r"corrispettivo\s+complessivo\s+pari\s+a\s*€\s*([0-9\.\,]+)", text, flags=re.I|re.S))
        if not importo:
            mensile = _parse_euro_num(_first(r"compenso\s+di\s*€\s*([0-9\.\,]+).*?\+\s*IVA\s+al\s+mese", text, flags=re.I|re.S))
            mesi_txt = _first(r"per\s+(\d{1,2})\s+mesi", text, flags=re.I|re.S)
            mesi = int(mesi_txt) if mesi_txt.isdigit() else 12
            importo = round(mensile * mesi, 2) if mensile else 0.0
        if not importo:
            importo = _parse_euro_num(_first(r"compenso\s+(?:complessivo|annuale).*?€\s*([0-9\.\,]+)", text, flags=re.I|re.S))

        # servizi Art.2: supporta sia "2.1 Titolo. testo" sia "2.1 Sviluppo e-commerce (Shopify) ..."
        for m in re.finditer(r"2\.\d+\s+(.+?)(?=\s+2\.\d+|\s+Art\.?\s*3|\s+ART\.?\s*3|$)", text, re.I|re.S):
            chunk = re.sub(r"\s+", " ", m.group(1)).strip(" .;:-")
            if not chunk:
                continue
            if "." in chunk[:120]:
                serv, desc = chunk.split(".", 1)
            else:
                serv, desc = chunk[:90], chunk
            servizi.append({"Servizio": serv.strip()[:120], "Descrizione": desc.strip()})

    # common
    if not rag and email_cliente:
        rag = email_cliente.split("@")[0].replace(".", " ").title()
    forma = "Impresa individuale" if re.search(r"impresa individuale|libero professionista", text, re.I) else ("S.r.l." if re.search(r"S\.?R\.?L\.?|SRL", text, re.I) else "")
    titolo = _first(r"(CONTRATTO\s+DI\s+[A-ZÀ-Ü\s]+?)(?:\s+sviluppo|\s+Cliente:|\s+PARTI|$)", text)
    titolo = titolo.title().replace(" Di ", " di ").replace(" E ", " e ") if titolo else "CONTRATTO DI CONSULENZA STRATEGICA E OPERATIVA"
    durata = int(_first(r"durata\s+di\s+(\d+)", text) or 12)
    return {
        "ragione_sociale": rag.strip(), "forma_giuridica": forma, "partita_iva": piva_cliente,
        "codice_fiscale": cf or piva_cliente, "sede_legale": sede.strip(), "pec": email_cliente if email_cliente.endswith(".pec.it") else "",
        "email": email_cliente, "telefono": tel, "legale_rappresentante": "",
        "titolo": titolo, "sottotitolo": "sviluppo digitale, visibilità estera, apertura mercati e richieste di appuntamento",
        "tipo_contratto": "Consulenza strategica e operativa", "data_firma": data_firma or date.today(),
        "data_decorrenza": data_firma or date.today(), "durata_mesi": durata, "importo_totale": importo,
        "iva_percentuale": 22.0, "modalita_pagamento": "Mensile" if (is_modulo_bebol or re.search(r"al\s+mese|mensile|ogni\s+mese", text, re.I)) else "Trimestrale",
        "foro_competente": _first(r"Foro\s+di\s+([A-ZÀ-Üa-zà-ü'’ ]+?)(?:,|\.|\s+salvo|$)", text) or "Napoli",
        "servizi": servizi,
        "raw_text": raw_text,
        "metodo_importo": "Somma rate modulo" if is_modulo_bebol and importo else ("Corrispettivo/compenso testuale" if importo else "Da verificare manualmente"),
        "qualita_lettura": "Buona" if rag and importo else ("Parziale: controllare i campi" if rag or importo else "Bassa: compilare/verificare manualmente"),
    }

# -----------------------------
# DB helpers
# -----------------------------
def conn():
    c = sqlite3.connect(DB_PATH)
    c.row_factory = sqlite3.Row
    return c

def execute(q: str, params: tuple = ()) -> int:
    with conn() as c:
        cur = c.execute(q, params)
        c.commit()
        return int(cur.lastrowid or 0)

def read_df(q: str, params: tuple = ()) -> pd.DataFrame:
    with conn() as c:
        return pd.read_sql_query(q, c, params=params)

def table_has_col(table: str, col: str) -> bool:
    with conn() as c:
        rows = c.execute(f"PRAGMA table_info({table})").fetchall()
    return col in [r["name"] for r in rows]

def add_col(table: str, col_def: str):
    col = col_def.split()[0]
    if not table_has_col(table, col):
        execute(f"ALTER TABLE {table} ADD COLUMN {col_def}")

def init_db():
    with conn() as c:
        c.executescript("""
        CREATE TABLE IF NOT EXISTS aziende (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            nome TEXT NOT NULL,
            forma_giuridica TEXT,
            piva TEXT,
            cf TEXT,
            sede TEXT,
            pec TEXT,
            codice_sdi TEXT,
            iban TEXT,
            telefono TEXT,
            email TEXT,
            logo_file TEXT,
            note TEXT,
            is_default INTEGER DEFAULT 0,
            created_at TEXT NOT NULL
        );
        CREATE TABLE IF NOT EXISTS clienti (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            ragione_sociale TEXT NOT NULL,
            forma_giuridica TEXT,
            partita_iva TEXT,
            codice_fiscale TEXT,
            rea TEXT,
            sede_legale TEXT,
            pec TEXT,
            codice_sdi TEXT,
            legale_rappresentante TEXT,
            telefono TEXT,
            email TEXT,
            codice_ateco TEXT,
            settore TEXT,
            stato_crm TEXT DEFAULT 'Attivo',
            note TEXT,
            created_at TEXT NOT NULL
        );
        CREATE TABLE IF NOT EXISTS templates_contratto (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            nome TEXT NOT NULL,
            descrizione TEXT,
            testo_base TEXT NOT NULL,
            attivo INTEGER DEFAULT 1,
            created_at TEXT NOT NULL
        );

        CREATE TABLE IF NOT EXISTS staff (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            username TEXT,
            nome TEXT NOT NULL,
            cognome TEXT,
            ruolo TEXT,
            email TEXT,
            telefono TEXT,
            password_hash TEXT,
            access_level TEXT DEFAULT 'Operativo Base',
            is_admin INTEGER DEFAULT 0,
            stato TEXT DEFAULT 'Attivo',
            note TEXT,
            created_at TEXT NOT NULL
        );
        CREATE TABLE IF NOT EXISTS contratti (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            azienda_id INTEGER,
            cliente_id INTEGER NOT NULL,
            template_id INTEGER,
            staff_id INTEGER,
            titolo TEXT NOT NULL,
            sottotitolo TEXT,
            tipo_contratto TEXT NOT NULL,
            data_firma TEXT,
            luogo_firma TEXT DEFAULT 'Napoli',
            data_decorrenza TEXT NOT NULL,
            data_scadenza TEXT NOT NULL,
            durata_mesi INTEGER NOT NULL,
            importo_totale REAL NOT NULL,
            iva_percentuale REAL NOT NULL,
            modalita_pagamento TEXT NOT NULL,
            foro_competente TEXT DEFAULT 'Napoli',
            stato TEXT NOT NULL,
            servizi_json TEXT,
            clausole_extra TEXT,
            note TEXT,
            file_docx TEXT,
            file_pdf TEXT,
            file_firmato TEXT,
            created_at TEXT NOT NULL,
            FOREIGN KEY(cliente_id) REFERENCES clienti(id),
            FOREIGN KEY(azienda_id) REFERENCES aziende(id)
        );
        CREATE TABLE IF NOT EXISTS pagamenti (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            contratto_id INTEGER NOT NULL,
            numero_rata INTEGER NOT NULL,
            data_scadenza TEXT NOT NULL,
            imponibile REAL NOT NULL,
            iva REAL NOT NULL,
            totale REAL NOT NULL,
            stato TEXT NOT NULL DEFAULT 'Da pagare',
            data_pagamento TEXT,
            note TEXT,
            created_at TEXT NOT NULL,
            FOREIGN KEY(contratto_id) REFERENCES contratti(id)
        );
        CREATE TABLE IF NOT EXISTS incassi_rate (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            pagamento_id INTEGER NOT NULL,
            importo_pagato REAL NOT NULL,
            tipo_movimento TEXT NOT NULL DEFAULT 'Acconto',
            data_pagamento TEXT NOT NULL,
            allegato_file TEXT,
            note TEXT,
            registrato_da_staff_id INTEGER,
            created_at TEXT NOT NULL,
            FOREIGN KEY(pagamento_id) REFERENCES pagamenti(id)
        );
        CREATE TABLE IF NOT EXISTS documenti (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            cliente_id INTEGER,
            contratto_id INTEGER,
            tipo TEXT,
            titolo TEXT NOT NULL,
            file_path TEXT,
            note TEXT,
            created_at TEXT NOT NULL
        );
        CREATE TABLE IF NOT EXISTS lavori (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            cliente_id INTEGER NOT NULL,
            contratto_id INTEGER,
            data_lavoro TEXT NOT NULL,
            tipo_lavoro TEXT,
            titolo TEXT NOT NULL,
            descrizione TEXT,
            stato TEXT,
            allegato_file TEXT,
            note_interne TEXT,
            created_at TEXT NOT NULL,
            FOREIGN KEY(cliente_id) REFERENCES clienti(id)
        );

        CREATE TABLE IF NOT EXISTS feedback_clienti (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            cliente_id INTEGER NOT NULL,
            contratto_id INTEGER,
            staff_id INTEGER,
            data_feedback TEXT NOT NULL,
            provenienza TEXT DEFAULT 'Cliente',
            valutazione INTEGER,
            testo_feedback TEXT,
            allegato_file TEXT,
            note TEXT,
            created_at TEXT NOT NULL
        );
        CREATE TABLE IF NOT EXISTS fatture (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            azienda_id INTEGER,
            cliente_id INTEGER NOT NULL,
            contratto_id INTEGER,
            pagamento_id INTEGER,
            lavoro_id INTEGER,
            numero TEXT NOT NULL,
            anno INTEGER NOT NULL,
            data_fattura TEXT NOT NULL,
            scadenza TEXT,
            descrizione TEXT,
            imponibile REAL NOT NULL,
            iva_percentuale REAL NOT NULL,
            iva REAL NOT NULL,
            totale REAL NOT NULL,
            stato TEXT DEFAULT 'Bozza',
            file_pdf TEXT,
            note TEXT,
            emessa_da_staff_id INTEGER,
            created_at TEXT NOT NULL
        );
        """)
        c.commit()

    # Migrazioni leggere per database già esistenti
    with conn() as c:
        cols = [r[1] for r in c.execute("PRAGMA table_info(contratti)").fetchall()]
        if "staff_id" not in cols:
            c.execute("ALTER TABLE contratti ADD COLUMN staff_id INTEGER")

        staff_cols = [r[1] for r in c.execute("PRAGMA table_info(staff)").fetchall()]
        if "username" not in staff_cols:
            c.execute("ALTER TABLE staff ADD COLUMN username TEXT")
        if "password_hash" not in staff_cols:
            c.execute("ALTER TABLE staff ADD COLUMN password_hash TEXT")
        if "access_level" not in staff_cols:
            c.execute("ALTER TABLE staff ADD COLUMN access_level TEXT DEFAULT 'Operativo Base'")
        if "is_admin" not in staff_cols:
            c.execute("ALTER TABLE staff ADD COLUMN is_admin INTEGER DEFAULT 0")

        incassi_cols = [r[1] for r in c.execute("PRAGMA table_info(incassi_rate)").fetchall()]
        if "registrato_da_staff_id" not in incassi_cols:
            c.execute("ALTER TABLE incassi_rate ADD COLUMN registrato_da_staff_id INTEGER")

        fatture_cols = [r[1] for r in c.execute("PRAGMA table_info(fatture)").fetchall()]
        if "emessa_da_staff_id" not in fatture_cols:
            c.execute("ALTER TABLE fatture ADD COLUMN emessa_da_staff_id INTEGER")
        c.commit()
    seed_defaults()

def seed_defaults():
    aziende = read_df("SELECT * FROM aziende")
    if aziende.empty:
        execute("""INSERT INTO aziende (nome, forma_giuridica, piva, cf, sede, pec, codice_sdi, is_default, created_at)
                   VALUES (?,?,?,?,?,?,?,?,?)""",
                ("BEBOL", "S.r.l.", "09814141215", "09814141215", "Via Benedetto Brin n. 63 – 80142 Napoli (NA)", "bebolsrl@pec.it", "", 1, now_iso()))
    tmpls = read_df("SELECT * FROM templates_contratto")
    if tmpls.empty:
        execute("INSERT INTO templates_contratto (nome, descrizione, testo_base, created_at) VALUES (?,?,?,?)",
                ("Contratto consulenza strategica e operativa BEBOL", "Base contrattuale con articoli 1-15, servizi dinamici e prospetto pagamenti.", BASE_CONTRATTO, now_iso()))
    # Garantisce credenziali minime per eventuali righe staff preesistenti senza login.
    with conn() as c:
        c.execute("UPDATE staff SET username = 'staff_' || id WHERE username IS NULL OR username = ''")
        c.execute("UPDATE staff SET password_hash = ? WHERE password_hash IS NULL OR password_hash = ''", (hash_password("staff123"),))
        c.commit()
    # Garantisce credenziali minime per eventuali righe staff preesistenti senza login.
    with conn() as c:
        c.execute("UPDATE staff SET username = 'staff_' || id WHERE username IS NULL OR username = ''")
        c.execute("UPDATE staff SET password_hash = ? WHERE password_hash IS NULL OR password_hash = ''", (hash_password("staff123"),))
        c.commit()

    staff_df = read_df("SELECT * FROM staff")
    # Garantisce SEMPRE l'esistenza dell'account admin/admin123, anche se il DB aveva già staff creati prima della funzione login.
    admin_df = read_df("SELECT * FROM staff WHERE LOWER(COALESCE(username,''))='admin' LIMIT 1")
    if admin_df.empty:
        execute("INSERT INTO staff (username,nome,cognome,ruolo,email,password_hash,access_level,is_admin,stato,note,created_at) VALUES (?,?,?,?,?,?,?,?,?,?,?)",
                ("admin", "Admin", "CUBE", "Amministratore", "admin@cube.local", hash_password("admin123"), "Amministratore", 1, "Attivo", "Utente amministratore predefinito. Cambiare password dopo il primo accesso.", now_iso()))
    else:
        admin_id = int(admin_df.iloc[0]["id"])
        execute("UPDATE staff SET password_hash=?, access_level='Amministratore', is_admin=1, stato='Attivo', nome=COALESCE(nome,'Admin'), cognome=COALESCE(cognome,'CUBE'), ruolo=COALESCE(ruolo,'Amministratore') WHERE id=?",
                (hash_password("admin123"), admin_id))

BASE_CONTRATTO = """
ART. 1 – OGGETTO GENERALE DELL’INCARICO
1.1 Il Cliente conferisce alla Società, che accetta, incarico di consulenza strategica e operativa finalizzato alla costruzione e al rafforzamento della presenza digitale del Cliente, alla predisposizione dei contenuti multilingua, allo sviluppo dei canali social professionali e all’apertura o impostazione di mercati esteri compatibili con il settore di attività.
1.2 L’incarico ha natura di prestazione di servizi e consulenza professionale, resa con autonomia organizzativa e tecnica, senza vincolo di subordinazione, senza rappresentanza e senza mandato a concludere contratti in nome e per conto del Cliente, salvo specifica autorizzazione scritta.

ART. 2 – SERVIZI COMPRESI
{{SERVIZI}}

ART. 3 – MODALITÀ DI ESECUZIONE E COLLABORAZIONE
3.1 La Società svolgerà l’incarico prevalentemente da remoto, mediante email, call, strumenti digitali, piattaforme collaborative e accessi tecnici conferiti dal Cliente. Eventuali incontri in presenza, trasferte o giornate operative presso la sede del Cliente dovranno essere concordati separatamente.
3.2 Il Cliente nominerà un referente operativo e fornirà tempestivamente materiali, fotografie, schede prodotto, listini, informazioni tecniche, traduzioni già disponibili, credenziali, autorizzazioni e approvazioni necessarie. Ritardi, incompletezze o mancate validazioni del Cliente potranno incidere sui tempi di consegna senza responsabilità della Società.
3.3 Gli accessi a siti, domini, hosting, social, marketplace, strumenti pubblicitari o piattaforme terze dovranno essere conferiti con permessi adeguati e sicuri.

ART. 4 – DURATA, SCADENZA E OPZIONE DI RINNOVO
4.1 Il presente contratto ha durata di {{DURATA_MESI}} mesi dalla data di sottoscrizione o dalla diversa data di decorrenza indicata dalle Parti per iscritto.
4.2 Alla scadenza il contratto cesserà automaticamente, salvo diverso accordo scritto o rinnovo concordato dalle Parti.
4.3 Le Parti potranno esercitare un’opzione di rinnovo mediante comunicazione scritta o PEC da trasmettere almeno 90 giorni prima della naturale scadenza.

ART. 5 – CORRISPETTIVO, IVA E PAGAMENTI ANTICIPATI
5.1 Per l’intero incarico il Cliente riconosce alla Società un corrispettivo complessivo pari a {{IMPORTO_TOTALE}} oltre IVA di legge.
5.2 Il pagamento avverrà in via anticipata rispetto al periodo di competenza secondo la modalità indicata nel prospetto pagamenti.
5.3 In caso di ritardo nei pagamenti, la Società potrà sospendere le attività dopo formale sollecito, fermo il diritto agli interessi moratori previsti dal D.Lgs. 231/2002, oltre spese di sollecito e recupero, se dovute.

ART. 6 – COSTI TERZI, ADV E SERVIZI NON INCLUSI
6.1 Il compenso remunera esclusivamente le attività professionali della Società. Restano a carico del Cliente, salvo diverso accordo scritto, costi di dominio, hosting, temi, app, plugin, marketplace, commissioni di vendita, software, traduzioni certificate, shooting fotografici, campagne pubblicitarie, consulenze legali/fiscali specialistiche, spedizioni, logistica, packaging e qualunque costo richiesto da piattaforme terze.
6.2 Le campagne pubblicitarie non sono incluse nel compenso salvo espressa indicazione tra i servizi compresi.

ART. 7 – OBBLIGHI DEL CLIENTE E CONFORMITÀ
7.1 Il Cliente garantisce titolarità o legittima disponibilità di marchi, immagini, testi, fotografie, schede tecniche, listini, descrizioni, claim commerciali e materiali forniti.
7.2 Il Cliente resta responsabile della conformità di prodotti, servizi, contenuti e informazioni alle normative applicabili e alle policy delle piattaforme.

ART. 8 – PRIVACY, DATI PERSONALI E COOKIE
8.1 Le Parti si impegnano a trattare i dati personali nel rispetto del Regolamento (UE) 2016/679, del D.Lgs. 196/2003 e normativa applicabile.
8.2 Il Cliente resta titolare del trattamento dei dati di clienti, utenti, lead e contatti raccolti tramite sito, social o piattaforme, salvo diverso accordo.

ART. 9 – PROPRIETÀ INTELLETTUALE E MATERIALI
9.1 Marchi, loghi, fotografie, testi originari, cataloghi e materiali commerciali del Cliente restano di titolarità del Cliente.
9.2 Metodologie, procedure, template, checklist, know-how, flussi operativi e impostazioni consulenziali predisposte dalla Società restano di titolarità della Società, salvo diversa pattuizione scritta.

ART. 10 – NATURA DEL SERVIZIO, RISULTATI E LIMITAZIONE DI RESPONSABILITÀ
10.1 L’attività della Società è obbligazione di mezzi e non di risultato: non sono garantiti specifici volumi di vendita, appuntamenti, posizionamenti, fatturati, ranking SEO, approvazioni marketplace o performance economiche.
10.2 La Società risponderà esclusivamente per dolo o colpa grave, nei limiti consentiti dalla legge.

ART. 11 – SOSPENSIONE, RECESSO E RISOLUZIONE
11.1 In caso di mancato pagamento di importi scaduti, la Società potrà sospendere le attività previa comunicazione scritta sino all’integrale saldo delle somme dovute.
11.2 Ciascuna Parte potrà recedere dal contratto per giusta causa con comunicazione scritta.
11.3 Ai sensi dell’art. 1456 c.c., il contratto potrà risolversi di diritto in caso di grave inadempimento.

ART. 12 – RISERVATEZZA
12.1 Le Parti si impegnano a mantenere riservate informazioni tecniche, commerciali, economiche, strategiche, credenziali, dati di accesso, listini, documenti, procedure e know-how appresi durante l’esecuzione del contratto.
12.2 L’obbligo di riservatezza permane per tutta la durata del contratto e per 24 mesi successivi alla cessazione.

ART. 13 – MANLEVA
13.1 Il Cliente si impegna a manlevare e tenere indenne la Società da pretese, reclami, contestazioni, sanzioni e richieste risarcitorie derivanti da prodotti, contenuti, claim, immagini, marchi, informazioni, vendite, resi, garanzie, fiscalità, privacy o policy di piattaforme riconducibili al Cliente, salvo dolo o colpa grave della Società.

ART. 14 – RIFERIMENTI NORMATIVI E ADEGUAMENTO LEGALE
14.1 Il presente contratto è interpretato secondo la legge italiana e secondo le norme del Codice Civile in materia di contratti, prestazione di servizi, correttezza professionale, inadempimento e responsabilità contrattuale.
14.2 Per attività digitali e online si richiamano, ove applicabili, il D.Lgs. 70/2003, il D.Lgs. 206/2005, il Regolamento (UE) 2016/679, il Regolamento (UE) 2022/2065 e il D.Lgs. 231/2002.

ART. 15 – LEGGE APPLICABILE, FORO COMPETENTE E DISPOSIZIONI FINALI
15.1 Il presente contratto costituisce l’intero accordo tra le Parti e sostituisce ogni precedente intesa avente il medesimo oggetto.
15.2 Per ogni controversia sarà competente in via esclusiva il Foro di {{FORO_COMPETENTE}}, salvo norme inderogabili.
15.3 Ai sensi degli artt. 1341 e 1342 c.c., il Cliente approva specificamente le clausole: artt. 3.2, 3.3, 4, 5, 6, 7.2, 8.2, 9.2, 10, 11, 12, 13, 14, 15.2.
"""

STATUS_COLORS = {
    "Attivo": ("#dcfce7", "#166534", "🟢"), "Firmato": ("#dbeafe", "#1d4ed8", "🔵"),
    "Bozza": ("#f1f5f9", "#334155", "⚪"), "Inviato": ("#e0f2fe", "#0369a1", "📨"),
    "Da pagare": ("#f1f5f9", "#334155", "⏳"), "Acconto": ("#ffedd5", "#c2410c", "🟠"),
    "Pagata": ("#dcfce7", "#166534", "✅"), "Scaduta": ("#fee2e2", "#b91c1c", "🔴"),
    "Scaduto": ("#fee2e2", "#b91c1c", "🔴"), "Sospeso": ("#fee2e2", "#7f1d1d", "⛔"),
    "Completato": ("#dcfce7", "#166534", "✅"), "In lavorazione": ("#e0f2fe", "#0369a1", "🔧"),
    "Consegnato": ("#dcfce7", "#166534", "📦"), "Bloccato": ("#fee2e2", "#b91c1c", "⛔"),
    "Pronta": ("#e0f2fe", "#0369a1", "📄"), "Stornata": ("#fee2e2", "#b91c1c", "↩️"),
}

def badge(status: str) -> str:
    bg, fg, ico = STATUS_COLORS.get(status, ("#f1f5f9", "#334155", "•"))
    return f"<span class='badge' style='background:{bg};color:{fg}'>{ico} {safe(status)}</span>"

# -----------------------------
# Business
# -----------------------------
def get_default_azienda_id() -> int:
    df = read_df("SELECT id FROM aziende WHERE is_default=1 LIMIT 1")
    if not df.empty:
        return int(df.iloc[0]["id"])
    df = read_df("SELECT id FROM aziende ORDER BY id LIMIT 1")
    return int(df.iloc[0]["id"]) if not df.empty else 0

def genera_rate(contratto_id: int, decorrenza: date, durata_mesi: int, imponibile: float, iva_pct: float, modalita: str):
    execute("DELETE FROM incassi_rate WHERE pagamento_id IN (SELECT id FROM pagamenti WHERE contratto_id=?)", (contratto_id,))
    execute("DELETE FROM pagamenti WHERE contratto_id=?", (contratto_id,))
    step = {"Mensile": 1, "Bimestrale": 2, "Trimestrale": 3, "Semestrale": 6, "Annuale": 12}.get(modalita, 1)
    n = max(1, durata_mesi // step)
    rata = round(imponibile / n, 2)
    for i in range(n):
        imp = round(imponibile - rata * (n - 1), 2) if i == n - 1 else rata
        iva = round(imp * iva_pct / 100, 2)
        execute("""INSERT INTO pagamenti (contratto_id, numero_rata, data_scadenza, imponibile, iva, totale, stato, created_at)
                   VALUES (?,?,?,?,?,?,?,?)""",
                (contratto_id, i + 1, add_months(decorrenza, i * step).isoformat(), imp, iva, round(imp + iva, 2), "Da pagare", now_iso()))


def get_or_create_cliente_from_import(dati: dict) -> int:
    """Crea o aggiorna il cliente importato dal contratto, così compare nel menu Clienti CRM."""
    rag = (dati.get("ragione_sociale") or "").strip() or "Cliente da verificare"
    piva = (dati.get("partita_iva") or "").strip()
    cf = (dati.get("codice_fiscale") or "").strip()
    # Cerca cliente esistente per P.IVA/C.F.; se lo trova, aggiorna i campi principali.
    existing = pd.DataFrame()
    if piva:
        existing = read_df("SELECT id FROM clienti WHERE partita_iva=? LIMIT 1", (piva,))
    if existing.empty and cf:
        existing = read_df("SELECT id FROM clienti WHERE codice_fiscale=? LIMIT 1", (cf,))
    if not existing.empty:
        cid = int(existing.iloc[0]["id"])
        execute("""UPDATE clienti SET
                    ragione_sociale=?, forma_giuridica=?, partita_iva=?, codice_fiscale=?,
                    sede_legale=?, pec=?, email=?, telefono=?, legale_rappresentante=?,
                    note=COALESCE(note,'')
                   WHERE id=?""",
                (rag, dati.get("forma_giuridica",""), piva, cf, dati.get("sede_legale",""),
                 dati.get("pec",""), dati.get("email",""), dati.get("telefono",""),
                 dati.get("legale_rappresentante",""), cid))
        return cid
    return execute("""INSERT INTO clienti (
                        ragione_sociale, forma_giuridica, partita_iva, codice_fiscale,
                        sede_legale, pec, email, telefono, legale_rappresentante,
                        stato_crm, note, created_at
                      ) VALUES (?,?,?,?,?,?,?,?,?,?,?,?)""",
                   (rag, dati.get("forma_giuridica",""), piva, cf, dati.get("sede_legale",""),
                    dati.get("pec",""), dati.get("email",""), dati.get("telefono",""),
                    dati.get("legale_rappresentante",""), "Attivo",
                    "Creato automaticamente da importazione contratto.", now_iso()))

def crea_contratto_da_import(dati: dict, file_name: str | None = None, file_bytes: bytes | None = None) -> int:
    """Salva davvero il contratto importato nel gestionale: cliente, contratto, rate, documenti, file Word/PDF."""
    azienda_id = get_default_azienda_id()
    cliente_id = get_or_create_cliente_from_import(dati)
    template_df = read_df("SELECT id FROM templates_contratto ORDER BY id LIMIT 1")
    template_id = int(template_df.iloc[0]["id"]) if not template_df.empty else None

    data_firma = parse_date(dati.get("data_firma"), date.today()) or date.today()
    decorrenza = parse_date(dati.get("data_decorrenza"), data_firma) or data_firma
    durata = int(dati.get("durata_mesi") or 12)
    scadenza = add_months(decorrenza, durata)
    importo = float(dati.get("importo_totale") or 0)
    iva_pct = float(dati.get("iva_percentuale") or 22.0)
    modalita = dati.get("modalita_pagamento") or "Mensile"
    servizi = dati.get("servizi") or []
    if isinstance(servizi, pd.DataFrame):
        servizi = servizi.to_dict("records")

    contratto_id = execute("""INSERT INTO contratti (
                    azienda_id, cliente_id, template_id, staff_id, titolo, sottotitolo, tipo_contratto,
                    data_firma, luogo_firma, data_decorrenza, data_scadenza, durata_mesi,
                    importo_totale, iva_percentuale, modalita_pagamento, foro_competente,
                    stato, servizi_json, clausole_extra, note, created_at
                ) VALUES (?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?)""",
                (azienda_id, cliente_id, template_id, None,
                 dati.get("titolo") or "CONTRATTO DI CONSULENZA STRATEGICA E OPERATIVA",
                 dati.get("sottotitolo") or "",
                 dati.get("tipo_contratto") or "Consulenza strategica e operativa",
                 data_firma.isoformat(), dati.get("luogo_firma") or "Napoli",
                 decorrenza.isoformat(), scadenza.isoformat(), durata,
                 importo, iva_pct, modalita, dati.get("foro_competente") or "Napoli",
                 "Importato", json.dumps(servizi, ensure_ascii=False), "",
                 "Contratto creato automaticamente tramite importazione PDF/foto.", now_iso()))

    # Genera rate: così il contratto appare anche in Pagamenti.
    genera_rate(contratto_id, decorrenza, durata, importo, iva_pct, modalita)

    # Allegato originale: così appare in Documenti e nella scheda contratto.
    if file_bytes and file_name:
        DOC_DIR.mkdir(parents=True, exist_ok=True)
        fname = f"{datetime.now().strftime('%Y%m%d_%H%M%S')}_contratto_importato_{slug(file_name)}"
        fpath = DOC_DIR / fname
        fpath.write_bytes(file_bytes)
        rel = str(fpath.relative_to(BASE_DIR))
        execute("INSERT INTO documenti (cliente_id,contratto_id,tipo,titolo,file_path,note,created_at) VALUES (?,?,?,?,?,?,?)",
                (cliente_id, contratto_id, "Contratto importato", "Contratto originale importato", rel,
                 "File originale caricato durante l'importazione.", now_iso()))
        execute("UPDATE contratti SET file_firmato=? WHERE id=?", (rel, contratto_id))

    # Genera Word/PDF interni basati sul template.
    try:
        generate_contract_files(contratto_id)
    except Exception:
        pass

    return contratto_id

def update_stato_pagamento(pagamento_id: int):
    df = read_df("SELECT totale FROM pagamenti WHERE id=?", (pagamento_id,))
    if df.empty: return
    totale = float(df.iloc[0]["totale"] or 0)
    inc = read_df("SELECT COALESCE(SUM(importo_pagato),0) pagato, MAX(data_pagamento) data_pagamento FROM incassi_rate WHERE pagamento_id=?", (pagamento_id,))
    pagato = float(inc.iloc[0]["pagato"] or 0)
    data_pag = inc.iloc[0]["data_pagamento"] if not inc.empty else None
    stato = "Da pagare" if pagato <= 0 else ("Pagata" if pagato + 0.01 >= totale else "Acconto")
    execute("UPDATE pagamenti SET stato=?, data_pagamento=? WHERE id=?", (stato, data_pag if pagato > 0 else None, pagamento_id))

def payments_df(where="", params=()):
    q = f"""
    SELECT p.*, cl.ragione_sociale cliente, c.titolo contratto,
           COALESCE(SUM(i.importo_pagato),0) pagato,
           ROUND(p.totale-COALESCE(SUM(i.importo_pagato),0),2) residuo,
           COUNT(i.id) movimenti
    FROM pagamenti p
    JOIN contratti c ON c.id=p.contratto_id
    JOIN clienti cl ON cl.id=c.cliente_id
    LEFT JOIN incassi_rate i ON i.pagamento_id=p.id
    {where}
    GROUP BY p.id
    ORDER BY p.data_scadenza ASC, p.id ASC
    """
    return read_df(q, params)

def servizi_to_text(servizi: list[dict]) -> str:
    if not servizi:
        return "2.1 Le attività specifiche saranno concordate dalle Parti in fase operativa e indicate nei documenti di progetto."
    lines = []
    for idx, s in enumerate(servizi, 1):
        titolo = s.get("Servizio") or s.get("servizio") or f"Servizio {idx}"
        desc = s.get("Descrizione") or s.get("descrizione") or "Attività inclusa nell'incarico."
        lines.append(f"2.{idx} {titolo}. {desc}")
    return "\n".join(lines)

def build_contract_text(azienda: dict, cliente: dict, contratto: dict, servizi: list[dict], template_text: str) -> str:
    txt = template_text.replace("{{SERVIZI}}", servizi_to_text(servizi))
    rep = {
        "{{DURATA_MESI}}": str(contratto.get("durata_mesi") or 12),
        "{{IMPORTO_TOTALE}}": money(contratto.get("importo_totale") or 0),
        "{{FORO_COMPETENTE}}": str(contratto.get("foro_competente") or "Napoli"),
    }
    for k,v in rep.items(): txt = txt.replace(k,v)
    return txt.strip()

def contract_records(contratto_id: int):
    cdf = read_df("SELECT * FROM contratti WHERE id=?", (contratto_id,))
    if cdf.empty: return None
    contratto = cdf.iloc[0].to_dict()
    azienda = read_df("SELECT * FROM aziende WHERE id=?", (contratto.get("azienda_id") or get_default_azienda_id(),)).iloc[0].to_dict()
    cliente = read_df("SELECT * FROM clienti WHERE id=?", (contratto["cliente_id"],)).iloc[0].to_dict()
    tdf = read_df("SELECT * FROM templates_contratto WHERE id=?", (contratto.get("template_id") or 1,))
    template = tdf.iloc[0].to_dict() if not tdf.empty else {"testo_base": BASE_CONTRATTO}
    try: servizi = json.loads(contratto.get("servizi_json") or "[]")
    except Exception: servizi = []
    return azienda, cliente, contratto, template, servizi

def generate_contract_files(contratto_id: int) -> tuple[str, str]:
    records = contract_records(contratto_id)
    if not records: raise ValueError("Contratto non trovato")
    azienda, cliente, contratto, template, servizi = records
    body = build_contract_text(azienda, cliente, contratto, servizi, template.get("testo_base") or BASE_CONTRATTO)
    fname = f"contratto_{contratto_id}_{slug(cliente['ragione_sociale'])}_{datetime.now().strftime('%Y%m%d_%H%M%S')}"
    docx_path = GENERATED_DIR / f"{fname}.docx"
    pdf_path = GENERATED_DIR / f"{fname}.pdf"

    # DOCX
    doc = Document()
    sec = doc.sections[0]
    sec.top_margin = Inches(0.65); sec.bottom_margin = Inches(0.65); sec.left_margin = Inches(0.7); sec.right_margin = Inches(0.7)
    if LOGO_PATH.exists():
        try: doc.add_picture(str(LOGO_PATH), width=Inches(1.7))
        except Exception: pass
    h = doc.add_heading(contratto.get("titolo") or "Contratto", 0)
    h.alignment = 1
    if contratto.get("sottotitolo"):
        p = doc.add_paragraph(contratto.get("sottotitolo")); p.alignment = 1
    doc.add_paragraph(f"Cliente: {cliente.get('ragione_sociale')}")
    doc.add_heading("PARTI CONTRAENTI", level=1)
    doc.add_paragraph(f"TRA\n{azienda.get('nome')}, {azienda.get('forma_giuridica') or ''}, con sede legale in {azienda.get('sede') or ''}, P.IVA/C.F. {azienda.get('piva') or azienda.get('cf') or ''}, PEC {azienda.get('pec') or ''}, di seguito denominata 'Società'.")
    doc.add_paragraph(f"E\n{cliente.get('ragione_sociale')}, {cliente.get('forma_giuridica') or ''}, con sede in {cliente.get('sede_legale') or ''}, P.IVA {cliente.get('partita_iva') or ''}, C.F. {cliente.get('codice_fiscale') or ''}, PEC {cliente.get('pec') or ''}, in persona di {cliente.get('legale_rappresentante') or ''}, di seguito denominata 'Cliente'.")
    doc.add_heading("PREMESSE", level=1)
    doc.add_paragraph("Le premesse e i dati delle Parti costituiscono parte integrante e sostanziale del presente contratto.")
    for block in body.split("\n\n"):
        if block.strip().startswith("ART."):
            lines = block.split("\n",1)
            doc.add_heading(lines[0].strip(), level=1)
            if len(lines)>1:
                for l in lines[1].split("\n"):
                    doc.add_paragraph(l.strip())
        else:
            doc.add_paragraph(block.strip())
    doc.add_heading("PROSPETTO MODALITÀ DI PAGAMENTO", level=1)
    pays = read_df("SELECT numero_rata, data_scadenza, imponibile, iva, totale FROM pagamenti WHERE contratto_id=? ORDER BY numero_rata", (contratto_id,))
    table = doc.add_table(rows=1, cols=5)
    for i,hdr in enumerate(["Rata", "Scadenza", "Imponibile", "IVA", "Totale"]): table.rows[0].cells[i].text = hdr
    for _,r in pays.iterrows():
        cells = table.add_row().cells
        cells[0].text = str(int(r["numero_rata"])); cells[1].text = str(r["data_scadenza"]); cells[2].text = money(r["imponibile"]); cells[3].text = money(r["iva"]); cells[4].text = money(r["totale"])
    doc.add_paragraph(f"LETTO, CONFERMATO E SOTTOSCRITTO\nLuogo e data: {contratto.get('luogo_firma') or ''}, {contratto.get('data_firma') or ''}")
    doc.add_paragraph(f"Per {azienda.get('nome')}\nFirma e Timbro: ____________________________")
    doc.add_paragraph(f"Per {cliente.get('ragione_sociale')}\nFirma e Timbro: ____________________________")
    doc.add_paragraph("Approvazione specifica ex artt. 1341 e 1342 c.c.\nPer il Cliente\nFirma e Timbro: ____________________________")
    doc.save(docx_path)

    # PDF
    styles = getSampleStyleSheet()
    styles.add(ParagraphStyle(name="CubeTitle", fontSize=17, leading=21, alignment=TA_CENTER, spaceAfter=12, textColor=colors.HexColor("#0c1d2f")))
    styles.add(ParagraphStyle(name="CubeH", fontSize=11, leading=14, spaceBefore=8, spaceAfter=5, textColor=colors.HexColor("#0c1d2f"), fontName="Helvetica-Bold"))
    styles.add(ParagraphStyle(name="CubeP", fontSize=8.8, leading=12, alignment=TA_LEFT))
    pdf = SimpleDocTemplate(str(pdf_path), pagesize=A4, rightMargin=1.5*cm, leftMargin=1.5*cm, topMargin=1.2*cm, bottomMargin=1.2*cm)
    story = [Paragraph(safe(contratto.get("titolo") or "Contratto"), styles["CubeTitle"]), Paragraph(f"Cliente: {safe(cliente.get('ragione_sociale'))}", styles["CubeP"]), Spacer(1,8)]
    story += [Paragraph("PARTI CONTRAENTI", styles["CubeH"]), Paragraph(safe(f"TRA {azienda.get('nome')}, con sede in {azienda.get('sede')}, P.IVA/C.F. {azienda.get('piva') or azienda.get('cf')}, PEC {azienda.get('pec')}"), styles["CubeP"]), Paragraph(safe(f"E {cliente.get('ragione_sociale')}, con sede in {cliente.get('sede_legale')}, P.IVA {cliente.get('partita_iva')}, C.F. {cliente.get('codice_fiscale')}, PEC {cliente.get('pec')}"), styles["CubeP"])]
    for block in body.split("\n\n"):
        if not block.strip(): continue
        if block.strip().startswith("ART."):
            lines = block.split("\n",1)
            story.append(Paragraph(safe(lines[0]), styles["CubeH"]))
            if len(lines)>1:
                for l in lines[1].split("\n"):
                    story.append(Paragraph(safe(l), styles["CubeP"]))
        else:
            story.append(Paragraph(safe(block), styles["CubeP"]))
    story.append(Paragraph("PROSPETTO MODALITÀ DI PAGAMENTO", styles["CubeH"]))
    data = [["Rata", "Scadenza", "Imponibile", "IVA", "Totale"]] + [[int(r["numero_rata"]), r["data_scadenza"], money(r["imponibile"]), money(r["iva"]), money(r["totale"])] for _,r in pays.iterrows()]
    tab = Table(data, colWidths=[1.5*cm, 3*cm, 3*cm, 3*cm, 3*cm])
    tab.setStyle(TableStyle([('BACKGROUND',(0,0),(-1,0), colors.HexColor('#0c1d2f')),('TEXTCOLOR',(0,0),(-1,0), colors.white),('GRID',(0,0),(-1,-1),0.25,colors.HexColor('#d9e2ec')),('FONT',(0,0),(-1,0),'Helvetica-Bold'),('FONTSIZE',(0,0),(-1,-1),8)]))
    story.append(tab); story.append(Spacer(1,12))
    story.append(Paragraph(safe(f"LETTO, CONFERMATO E SOTTOSCRITTO - {contratto.get('luogo_firma')}, {contratto.get('data_firma')}"), styles["CubeP"]))
    story.append(Spacer(1,18)); story.append(Paragraph(safe(f"Per {azienda.get('nome')} ____________________________    Per {cliente.get('ragione_sociale')} ____________________________"), styles["CubeP"]))
    pdf.build(story)
    rel_docx = str(docx_path.relative_to(BASE_DIR)); rel_pdf = str(pdf_path.relative_to(BASE_DIR))
    execute("UPDATE contratti SET file_docx=?, file_pdf=? WHERE id=?", (rel_docx, rel_pdf, contratto_id))
    execute("INSERT INTO documenti (cliente_id, contratto_id, tipo, titolo, file_path, note, created_at) VALUES (?,?,?,?,?,?,?)", (contratto["cliente_id"], contratto_id, "Contratto", "Contratto generato DOCX", rel_docx, "Generato automaticamente", now_iso()))
    execute("INSERT INTO documenti (cliente_id, contratto_id, tipo, titolo, file_path, note, created_at) VALUES (?,?,?,?,?,?,?)", (contratto["cliente_id"], contratto_id, "Contratto", "Contratto generato PDF", rel_pdf, "Generato automaticamente", now_iso()))
    return rel_docx, rel_pdf

def generate_invoice_pdf(fattura_id: int) -> str:
    fdf = read_df("SELECT f.*, cl.ragione_sociale cliente, cl.sede_legale, cl.partita_iva, cl.codice_fiscale, cl.pec, a.nome azienda, a.piva azienda_piva, a.sede azienda_sede, a.pec azienda_pec FROM fatture f JOIN clienti cl ON cl.id=f.cliente_id LEFT JOIN aziende a ON a.id=f.azienda_id WHERE f.id=?", (fattura_id,))
    if fdf.empty: raise ValueError("Fattura non trovata")
    f = fdf.iloc[0]
    fname = f"fattura_{slug(f['numero'])}_{slug(f['cliente'])}.pdf"
    path = INV_DIR / fname
    styles = getSampleStyleSheet()
    styles.add(ParagraphStyle(name="TitleInv", fontSize=18, leading=22, alignment=TA_CENTER, textColor=colors.HexColor("#0c1d2f")))
    pdf = SimpleDocTemplate(str(path), pagesize=A4, rightMargin=1.8*cm, leftMargin=1.8*cm, topMargin=1.6*cm, bottomMargin=1.4*cm)
    story = [Paragraph("FATTURA DI CORTESIA", styles["TitleInv"]), Spacer(1,12)]
    story.append(Paragraph(safe(f"Emittente: {f['azienda']} - P.IVA {f['azienda_piva']} - {f['azienda_sede']} - PEC {f['azienda_pec']}"), styles["Normal"]))
    story.append(Paragraph(safe(f"Cliente: {f['cliente']} - P.IVA {f['partita_iva']} - C.F. {f['codice_fiscale']} - {f['sede_legale']} - PEC {f['pec']}"), styles["Normal"]))
    story.append(Spacer(1,12)); story.append(Paragraph(safe(f"Numero: {f['numero']} del {f['data_fattura']} - Stato: {f['stato']}"), styles["Normal"]))
    data = [["Descrizione", "Imponibile", "IVA %", "IVA", "Totale"], [f["descrizione"], money(f["imponibile"]), f["iva_percentuale"], money(f["iva"]), money(f["totale"])] ]
    tab = Table(data, colWidths=[7*cm, 3*cm, 2*cm, 3*cm, 3*cm])
    tab.setStyle(TableStyle([('BACKGROUND',(0,0),(-1,0), colors.HexColor('#0c1d2f')),('TEXTCOLOR',(0,0),(-1,0), colors.white),('GRID',(0,0),(-1,-1),0.25,colors.grey),('VALIGN',(0,0),(-1,-1),'TOP')]))
    story.append(tab); story.append(Spacer(1,18)); story.append(Paragraph("Documento interno/di cortesia. Non sostituisce la fattura elettronica XML trasmessa allo SdI.", styles["Italic"]))
    pdf.build(story)
    rel = str(path.relative_to(BASE_DIR))
    execute("UPDATE fatture SET file_pdf=? WHERE id=?", (rel, fattura_id))
    return rel

# -----------------------------
# UI
# -----------------------------
def css():
    st.markdown("""
    <style>
    .stApp { background: linear-gradient(180deg,#f5f8fc 0%,#edf3f9 100%); }
    [data-testid="stSidebar"] { background: linear-gradient(180deg,#09192b 0%,#0e2742 100%); }
    [data-testid="stSidebar"] * { color:white; }
    .hero { background: linear-gradient(135deg,rgba(12,29,47,.06),rgba(15,109,208,.10)); border:1px solid #dbe6f3; border-radius:22px; padding:20px; margin-bottom:12px; }
    .hero h1 { margin:0; color:#0c1d2f; font-size:2.2rem; }
    .hero p { margin:.35rem 0 0 0; color:#64748b; }
    .badge { display:inline-block; padding:5px 10px; border-radius:999px; font-size:.84rem; font-weight:800; }
    .section { display:flex; gap:12px; align-items:center; margin:14px 0 8px; }
    .section .ico { width:42px;height:42px;border-radius:14px;display:flex;align-items:center;justify-content:center;background:linear-gradient(135deg,#0f6dd0,#0c1d2f);color:white;box-shadow:0 8px 18px rgba(15,109,208,.22); }
    .section .ttl { font-size:1.65rem; font-weight:900; color:#17263c; }
    .section .cap { color:#66778d; font-size:.94rem; margin-top:2px; }
    .card { background:white; border:1px solid #dfe8f2; border-radius:18px; padding:15px 17px; box-shadow:0 8px 22px rgba(12,29,47,.06); min-height:100px; }
    .card .icon { font-size:1.25rem; margin-bottom:5px; }.card .label{font-size:.85rem;color:#64748b;font-weight:800}.card .value{font-size:1.05rem;color:#17263c;font-weight:900;word-break:break-word}
    div[data-testid="stMetric"] { background:white; padding:16px; border-radius:18px; border:1px solid #dfe8f2; box-shadow:0 8px 22px rgba(12,29,47,.06); }
    div.stButton > button, div.stDownloadButton > button, div[data-testid="stFormSubmitButton"] button { border-radius:12px!important; background:linear-gradient(135deg,#0f6dd0,#0c5bb0)!important; color:white!important; border:0!important; font-weight:800!important; }
    div[data-testid="stDataFrame"] { background:white; border-radius:14px; border:1px solid #dfe8f2; padding:4px; }
    .smallnote { color:#64748b; font-size:.9rem; }
    </style>
    """, unsafe_allow_html=True)

def header():
    col1, col2 = st.columns([1,4])
    with col1:
        if LOGO_PATH.exists(): st.image(str(LOGO_PATH), use_container_width=True)
    with col2:
        st.markdown("<div class='hero'><span class='badge' style='background:#e9f2ff;color:#0f53a5'>📘 Streamlit Pro rifattibile</span><h1>CUBE Management Contract</h1><p>Contratti, CRM, documenti, lavori, pagamenti e fatture interne.</p></div>", unsafe_allow_html=True)

def section(icon,title,cap=""):
    st.markdown(f"<div class='section'><div class='ico'>{icon}</div><div><div class='ttl'>{safe(title)}</div>{('<div class=cap>'+safe(cap)+'</div>') if cap else ''}</div></div>", unsafe_allow_html=True)

def card(icon,label,value):
    st.markdown(f"<div class='card'><div class='icon'>{icon}</div><div class='label'>{safe(label)}</div><div class='value'>{safe(value)}</div></div>", unsafe_allow_html=True)




def staff_access_level() -> str:
    user = current_staff_user()
    if not user:
        return ""
    if int(user.get("is_admin") or 0) == 1:
        return "Amministratore"
    return str(user.get("access_level") or "Operativo Base")

def can_see_all_clients() -> bool:
    return staff_access_level() in ["Gestione Finanziaria", "Amministratore"]

def can_see_contracts_operational() -> bool:
    return staff_access_level() in ["Manager Operativo", "Amministratore"]

def can_see_money() -> bool:
    return staff_access_level() in ["Gestione Finanziaria", "Amministratore"]

def can_manage_finance() -> bool:
    return staff_access_level() in ["Gestione Finanziaria", "Amministratore"]

def user_is_admin() -> bool:
    user = current_staff_user()
    return bool(user and (int(user.get("is_admin") or 0) == 1 or str(user.get("access_level") or "") == "Amministratore"))

def assigned_contracts_df() -> pd.DataFrame:
    user = current_staff_user()
    if not user:
        return pd.DataFrame()
    if user_is_admin():
        return read_df("SELECT * FROM contratti ORDER BY id DESC")
    if staff_access_level() == "Gestione Finanziaria":
        return read_df("SELECT * FROM contratti ORDER BY id DESC")
    return read_df("SELECT * FROM contratti WHERE staff_id=? ORDER BY id DESC", (int(user["id"]),))

def assigned_clients_df() -> pd.DataFrame:
    """Clienti visibili dall'utente loggato in base al livello staff."""
    user = current_staff_user()
    if not user:
        return pd.DataFrame()
    if can_see_all_clients(): 
        return read_df("SELECT DISTINCT id, ragione_sociale, partita_iva FROM clienti ORDER BY ragione_sociale")
    return read_df("""
        SELECT DISTINCT cl.id, cl.ragione_sociale, cl.partita_iva
        FROM clienti cl
        JOIN contratti c ON c.cliente_id = cl.id
        WHERE c.staff_id = ?
        ORDER BY cl.ragione_sociale
    """, (int(user["id"]),))

def require_admin():
    if not user_is_admin():
        st.error("Accesso riservato all'amministratore.")
        st.stop()

def login_screen() -> bool:
    """Schermata di login staff. Ritorna True se autenticato."""
    if st.session_state.get("staff_user_id"):
        return True

    st.markdown("<br><br>", unsafe_allow_html=True)
    if LOGO_PATH.exists():
        c_logo, c_box = st.columns([1, 2])
        with c_logo:
            st.image(str(LOGO_PATH), use_container_width=True)
    st.title("CUBE Management Contract")
    st.subheader("Login staff")

    with st.form("login_staff_form"):
        username = st.text_input("Username o email")
        password = st.text_input("Password", type="password")
        submit = st.form_submit_button("🔐 Accedi")
    if submit:
        q = """
            SELECT * FROM staff
            WHERE stato='Attivo'
              AND (LOWER(COALESCE(username,'')) = LOWER(?) OR LOWER(COALESCE(email,'')) = LOWER(?))
            LIMIT 1
        """
        df = read_df(q, (username.strip(), username.strip()))
        if df.empty:
            st.error("Utente non trovato o non attivo.")
        else:
            row = df.iloc[0].to_dict()
            if verify_password(password, row.get("password_hash")):
                st.session_state["staff_user_id"] = int(row["id"])
                st.session_state["staff_username"] = row.get("username") or row.get("email") or row.get("nome")
                st.success("Accesso effettuato.")
                st.rerun()
            else:
                st.error("Password non corretta.")

    st.info("Primo accesso amministratore predefinito: username `admin`, password `admin123`. Cambiala dalla sezione Admin / Staff.")
    return False

def sidebar():
    user = current_staff_user()
    if LOGO_PATH.exists():
        st.sidebar.image(str(LOGO_PATH), use_container_width=True)
    if user:
        nome_user = f"{user.get('nome') or ''} {user.get('cognome') or ''}".strip()
        livello = staff_access_level()
        st.sidebar.markdown(f"**👤 {nome_user or user.get('username','Staff')}**")
        st.sidebar.caption(f"{user.get('ruolo') or 'Membro staff'} · {livello}")
        if st.sidebar.button("🚪 Esci"):
            st.session_state.clear()
            st.rerun()

    if user_is_admin():
        opts = {
            "🏠 Dashboard":"Dashboard",
            "📥 Importa contratto":"Importa contratto",
            "🏢 Aziende":"Aziende",
            "👥 Clienti CRM":"Clienti CRM",
            "📝 Crea nuovo contratto":"Crea nuovo contratto",
            "📚 Contratti":"Contratti",
            "💶 Pagamenti":"Pagamenti",
            "🛠️ Lavori":"Lavori",
            "📎 Documenti":"Documenti",
            "🧾 Fatture":"Fatture",
            "📄 Template contratti":"Template",
            "⚙️ Impostazioni":"Impostazioni",
            "👤 Admin / Staff":"Admin",
        }
    elif staff_access_level() == "Gestione Finanziaria":
        # Gestione finanziaria: vede clienti, pagamenti, fatture e documenti. Non vede lavori operativi né admin.
        opts = {
            "🏠 Dashboard finanziaria":"Dashboard",
            "👥 Clienti CRM":"Clienti CRM",
            "💶 Pagamenti":"Pagamenti",
            "🧾 Fatture":"Fatture",
            "📎 Documenti":"Documenti",
        }
    elif staff_access_level() == "Manager Operativo":
        # Manager operativo: vede gestione operativa più ampia, ma non denaro, rate, fatture o admin.
        opts = {
            "🏠 Dashboard operativa":"Dashboard",
            "👥 Clienti CRM":"Clienti CRM",
            "📚 Contratti operativi":"Contratti",
            "🛠️ Lavori":"Lavori",
            "📎 Documenti":"Documenti",
        }
    elif staff_access_level() == "Operativo Avanzato":
        # Operativo avanzato: vede tutti i clienti/lavori/documenti, senza contratti economici.
        opts = {
            "🏠 Dashboard operativa":"Dashboard",
            "👥 Clienti CRM":"Clienti CRM",
            "🛠️ Lavori":"Lavori",
            "📎 Documenti":"Documenti",
        }
    else:
        # Operativo base: solo clienti assegnati, lavori e documenti.
        opts = {
            "🏠 Dashboard operativa":"Dashboard",
            "👥 Clienti assegnati":"Clienti CRM",
            "🛠️ Lavori":"Lavori",
            "📎 Documenti":"Documenti",
        }

    labels = list(opts.keys())
    requested = st.session_state.pop("nav_to", None)
    if requested in labels:
        st.session_state["main_menu"] = requested
    current = st.session_state.get("main_menu", labels[0])
    index = labels.index(current) if current in labels else 0
    selected = st.sidebar.radio("Menu", labels, index=index, key="main_menu", label_visibility="collapsed")
    return opts[selected]


def company_select(label="Azienda emittente", key="azienda"):
    df = read_df("SELECT id,nome,piva FROM aziende ORDER BY is_default DESC, nome")
    opts = {f"{r['nome']} · P.IVA {r['piva'] or '-'} · ID {r['id']}": int(r['id']) for _,r in df.iterrows()}
    return st.selectbox(label, list(opts.keys()), key=key), opts

def client_select(label="Cliente", key="cliente"):
    df = assigned_clients_df()
    if df.empty:
        return None, {}
    opts = {f"{r['ragione_sociale']} · P.IVA {r['partita_iva'] or '-'} · ID {r['id']}": int(r['id']) for _,r in df.iterrows()}
    return st.selectbox(label, list(opts.keys()), key=key), opts


def staff_select(label="Membro staff", key="staff", include_none=True):
    df = read_df("SELECT id,nome,cognome,ruolo,email,stato FROM staff ORDER BY nome,cognome")
    opts = {}
    if include_none:
        opts["Non assegnato"] = None
    for _, r in df.iterrows():
        nome = f"{r['nome'] or ''} {r['cognome'] or ''}".strip()
        ruolo = r['ruolo'] or '-'
        stato = r['stato'] or '-'
        opts[f"{nome} · {ruolo} · {stato} · ID {int(r['id'])}"] = int(r['id'])
    if not opts:
        opts["Non assegnato"] = None
    return st.selectbox(label, list(opts.keys()), key=key), opts

def status_html_series(s):
    return s.apply(lambda x: re.sub('<.*?>','',badge(str(x))))

def applica_import_nei_campi(dati: dict):
    """Copia i dati letti dal PDF/foto nei campi manuali della scheda Crea nuovo contratto."""
    st.session_state["new_cliente_mode"] = "Nuovo cliente manuale"
    st.session_state["new_rag"] = dati.get("ragione_sociale", "")
    st.session_state["new_piva"] = dati.get("partita_iva", "")
    st.session_state["new_cf"] = dati.get("codice_fiscale", "")
    st.session_state["new_sede"] = dati.get("sede_legale", "")
    st.session_state["new_pec"] = dati.get("pec", "")
    st.session_state["new_email"] = dati.get("email", "")
    st.session_state["new_tel"] = dati.get("telefono", "")
    st.session_state["new_leg"] = dati.get("legale_rappresentante", "")
    st.session_state["new_titolo"] = dati.get("titolo", "CONTRATTO DI CONSULENZA STRATEGICA E OPERATIVA")
    st.session_state["new_sottotitolo"] = dati.get("sottotitolo", "sviluppo digitale, visibilità estera, apertura mercati e richieste di appuntamento")
    st.session_state["new_tipo"] = dati.get("tipo_contratto", "Consulenza strategica e operativa")
    st.session_state["new_data_firma"] = parse_date(dati.get("data_firma"), date.today())
    st.session_state["new_luogo"] = dati.get("luogo_firma", "Napoli")
    st.session_state["new_decorrenza"] = parse_date(dati.get("data_decorrenza"), date.today())
    st.session_state["new_durata"] = int(dati.get("durata_mesi", 12) or 12)
    st.session_state["new_importo"] = float(dati.get("importo_totale", 6000.0) or 0.0)
    st.session_state["new_iva_pct"] = float(dati.get("iva_percentuale", 22.0) or 22.0)
    st.session_state["new_modalita"] = dati.get("modalita_pagamento", "Trimestrale")
    st.session_state["new_foro"] = dati.get("foro_competente", "Napoli")
    st.session_state["new_note"] = "Creato da importazione contratto."
    if dati.get("servizi"):
        st.session_state.new_services = pd.DataFrame(dati.get("servizi"))
    st.session_state["contratto_importato_nei_campi"] = True


def blocco_importa_contratto(key_prefix: str = "global"):
    """Blocco riutilizzabile: carica PDF/foto, legge dati, importa nei campi compilabili."""
    uploaded_contract = st.file_uploader(
        "Carica PDF o foto contratto",
        type=["pdf", "png", "jpg", "jpeg", "webp"],
        key=f"{key_prefix}_upload_contract_import",
    )
    cimp1, cimp2 = st.columns([1, 2])
    with cimp1:
        leggi = st.button("📖 Leggi PDF/foto ed estrai i dati", use_container_width=True, key=f"{key_prefix}_leggi_import")
    with cimp2:
        st.info("Dopo la lettura puoi importare i dati nei campi compilabili e poi completarli/modificarli.")

    if leggi:
        if uploaded_contract is None:
            st.error("Carica prima un PDF o una foto del contratto.")
        else:
            try:
                raw = estrai_testo_da_upload_contratto(uploaded_contract)
                dati = estrai_dati_contratto_da_testo(raw)
                st.session_state.import_contract_data = dati
                st.session_state.import_contract_file_name = uploaded_contract.name
                st.session_state.import_contract_file_bytes = get_upload_bytes(uploaded_contract)
                if dati.get("servizi"):
                    st.session_state.new_services = pd.DataFrame(dati.get("servizi"))
                st.success("Dati estratti. Controlla il riepilogo e poi clicca 'Importa contratto nei campi compilabili'.")
                sc1, sc2, sc3, sc4 = st.columns(4)
                with sc1:
                    card("👤", "Cliente rilevato", dati.get("ragione_sociale") or "Da verificare")
                with sc2:
                    card("🧾", "P.IVA / C.F.", dati.get("partita_iva") or dati.get("codice_fiscale") or "Da verificare")
                with sc3:
                    card("📍", "Sede", dati.get("sede_legale") or "Da verificare")
                with sc4:
                    card("💶", "Importo", money(dati.get("importo_totale")) if dati.get("importo_totale") else "Da verificare")
                st.caption(f"Qualità lettura: {dati.get('qualita_lettura','Da verificare')} · Metodo importo: {dati.get('metodo_importo','Da verificare')}")
                st.info("Prima di salvare/importare puoi sempre correggere importo, servizi e anagrafica nella scheda Contratto. Per foto o scansioni inclinate, controlla sempre il riepilogo.")
                st.caption(f"Qualità lettura: {dati.get('qualita_lettura','Da verificare')} · Metodo importo: {dati.get('metodo_importo','Da verificare')}")
                st.info("Prima di salvare/importare puoi sempre correggere importo, servizi e anagrafica nella scheda Contratto. Per foto o scansioni inclinate, controlla sempre il riepilogo.")
                if dati.get("servizi"):
                    st.markdown("**Servizi rilevati dal contratto:**")
                    st.dataframe(pd.DataFrame(dati.get("servizi")), use_container_width=True, hide_index=True)
                with st.expander("Anteprima testo letto", expanded=False):
                    st.text_area("Testo estratto", clean_import_preview_text(raw)[:15000], height=260, key=f"{key_prefix}_raw_preview")
            except Exception as exc:
                st.error(str(exc))

    if st.session_state.get("import_contract_data"):
        st.divider()
        section("✅", "Importa contratto nei campi", "Trasferisce i dati letti nella scheda di creazione contratto, dove puoi modificarli.")
        st.warning("Se vuoi solo correggere i campi prima della creazione, usa il primo pulsante. Se invece vuoi registrare subito il contratto nel gestionale, usa il pulsante blu centrale.")
        col_importa, col_salva_gestionale, col_pulisci = st.columns([1, 1.35, 1])
        with col_importa:
            if st.button("✅ Porta nei campi compilabili", use_container_width=True, key=f"{key_prefix}_btn_importa_contratto"):
                applica_import_nei_campi(st.session_state.import_contract_data)
                st.success("Dati inseriti nei campi compilabili.")
        with col_salva_gestionale:
            if st.button("💾 Importa contratto nel gestionale", use_container_width=True, key=f"{key_prefix}_btn_salva_contratto_gestionale"):
                try:
                    contratto_id = crea_contratto_da_import(
                        st.session_state.import_contract_data,
                        st.session_state.get("import_contract_file_name"),
                        st.session_state.get("import_contract_file_bytes"),
                    )
                    st.session_state["ultimo_contratto_importato_id"] = contratto_id
                    st.success(f"Contratto importato nel gestionale. ID contratto: {contratto_id}. Ora è visibile in Clienti CRM, Contratti, Pagamenti e Documenti.")
                except Exception as exc:
                    st.error(f"Non sono riuscito a salvare il contratto nel gestionale: {exc}")
        with col_pulisci:
            if st.button("🧹 Pulisci importazione", use_container_width=True, key=f"{key_prefix}_btn_pulisci_import"):
                for k in ["import_contract_data", "import_contract_file_name", "import_contract_file_bytes", "contratto_importato_nei_campi", "ultimo_contratto_importato_id"]:
                    if k in st.session_state:
                        del st.session_state[k]
                st.session_state.new_services = pd.DataFrame([{"Servizio": "", "Descrizione": ""}])
                st.rerun()
        if st.session_state.get("contratto_importato_nei_campi"):
            st.success("Dati inseriti nei campi. Apri '📝 Crea nuovo contratto' e completa/verifica prima della generazione.")
            if st.button("➡️ Vai a Crea nuovo contratto", use_container_width=True, key=f"{key_prefix}_vai_crea"):
                st.session_state["nav_to"] = "📝 Crea nuovo contratto"
                st.rerun()
        if st.session_state.get("ultimo_contratto_importato_id"):
            cid = int(st.session_state["ultimo_contratto_importato_id"])
            cgo1, cgo2, cgo3 = st.columns(3)
            with cgo1:
                if st.button("📚 Apri Contratti", use_container_width=True, key=f"{key_prefix}_vai_contratti"):
                    st.session_state["nav_to"] = "📚 Contratti"
                    st.rerun()
            with cgo2:
                if st.button("💶 Apri Pagamenti", use_container_width=True, key=f"{key_prefix}_vai_pagamenti"):
                    st.session_state["nav_to"] = "💶 Pagamenti"
                    st.rerun()
            with cgo3:
                if st.button("👥 Apri Clienti CRM", use_container_width=True, key=f"{key_prefix}_vai_clienti"):
                    st.session_state["nav_to"] = "👥 Clienti CRM"
                    st.rerun()

# -----------------------------
# Pages
# -----------------------------
def page_dashboard():
    header(); section("🏠","Dashboard","Panoramica generale del gestionale rifattibile.")
    if staff_access_level() == "Gestione Finanziaria":
        section("🏠","Dashboard finanziaria","Area riservata alla gestione finanziaria: incassi, rate, fatture e documenti.")
        clienti_fin = assigned_clients_df()
        p_fin = payments_df()
        fatt_fin = read_df("SELECT * FROM fatture")
        totale_rate = float(p_fin["totale"].fillna(0).sum()) if not p_fin.empty else 0.0
        incassato = float(p_fin["pagato"].fillna(0).sum()) if not p_fin.empty else 0.0
        residuo = float(p_fin["residuo"].fillna(0).clip(lower=0).sum()) if not p_fin.empty else 0.0
        c1,c2,c3,c4 = st.columns(4)
        c1.metric("Clienti visibili", len(clienti_fin))
        c2.metric("Totale rate", money(totale_rate))
        c3.metric("Incassato", money(incassato))
        c4.metric("Residuo", money(residuo))
        c5,c6 = st.columns(2)
        c5.metric("Fatture", len(fatt_fin))
        c6.metric("Rate aperte", len(p_fin[p_fin["stato"].isin(["Da pagare","Acconto","Sollecitata"])]) if not p_fin.empty else 0)
        section("💶","Rate da gestire")
        if p_fin.empty:
            st.info("Nessuna rata.")
        else:
            v = p_fin[p_fin["stato"]!="Pagata"].head(20)[["cliente","contratto","numero_rata","data_scadenza","imponibile","iva","totale","pagato","residuo","stato"]].copy()
            for col in ["imponibile","iva","totale","pagato","residuo"]:
                v[col]=v[col].apply(money)
            st.dataframe(v, use_container_width=True, hide_index=True)
        return

    if not user_is_admin():
        user = current_staff_user()
        livello = staff_access_level()

        # La parte operativa non mostra clienti registrati generali: mostra solo i clienti assegnati allo staff.
        if livello == "Gestione Finanziaria":
            # Questo ramo resta finanziario e non operativo.
            section("🏠","Dashboard finanziaria","Area riservata alla gestione finanziaria: incassi, rate, fatture e documenti.")
            clienti_fin = assigned_clients_df()
            p_fin = payments_df()
            fatt_fin = read_df("SELECT * FROM fatture")
            totale_rate = float(p_fin["totale"].fillna(0).sum()) if not p_fin.empty else 0.0
            incassato = float(p_fin["pagato"].fillna(0).sum()) if not p_fin.empty else 0.0
            residuo = float(p_fin["residuo"].fillna(0).clip(lower=0).sum()) if not p_fin.empty else 0.0
            c1,c2,c3,c4 = st.columns(4)
            c1.metric("Clienti visibili", len(clienti_fin))
            c2.metric("Totale rate", money(totale_rate))
            c3.metric("Incassato", money(incassato))
            c4.metric("Residuo", money(residuo))
            c5,c6 = st.columns(2)
            c5.metric("Fatture", len(fatt_fin))
            c6.metric("Rate aperte", len(p_fin[p_fin["stato"].isin(["Da pagare","Acconto","Sollecitata"])]) if not p_fin.empty else 0)
            section("💶","Rate da gestire")
            if p_fin.empty:
                st.info("Nessuna rata.")
            else:
                v = p_fin[p_fin["stato"]!="Pagata"].head(20)[["cliente","contratto","numero_rata","data_scadenza","imponibile","iva","totale","pagato","residuo","stato"]].copy()
                for col in ["imponibile","iva","totale","pagato","residuo"]:
                    v[col]=v[col].apply(money)
                st.dataframe(v, use_container_width=True, hide_index=True)
            return

        clienti_staff = assigned_clients_df()
        contratti_staff = assigned_contracts_df()
        lavori_staff = read_df("""
            SELECT DISTINCT l.id, cl.ragione_sociale cliente, l.data_lavoro, l.tipo_lavoro, l.titolo, l.stato, l.allegato_file
            FROM lavori l
            JOIN clienti cl ON cl.id = l.cliente_id
            JOIN contratti c ON c.cliente_id = cl.id
            WHERE c.staff_id = ?
            ORDER BY l.data_lavoro DESC, l.id DESC
        """, (int(user["id"]),))
        feedback_staff = read_df("""
            SELECT DISTINCT f.id, cl.ragione_sociale cliente, f.data_feedback, f.provenienza, f.valutazione, f.testo_feedback
            FROM feedback_clienti f
            JOIN clienti cl ON cl.id = f.cliente_id
            JOIN contratti c ON c.cliente_id = cl.id
            WHERE c.staff_id = ?
            ORDER BY f.data_feedback DESC, f.id DESC
        """, (int(user["id"]),))

        section("🏠","Dashboard operativa","Area riservata allo staff: clienti assegnati, lavori, documenti e feedback. I dati economici non sono visibili.")
        c1,c2,c3 = st.columns(3)
        c1.metric("N. clienti assegnati", len(clienti_staff))
        c2.metric("Totale lavori registrati", len(lavori_staff))
        c3.metric("Feedback ricevuti", len(feedback_staff))

        st.divider()
        section("👥","Clienti assegnati")
        if clienti_staff.empty:
            st.info("Nessun cliente assegnato.")
        else:
            st.dataframe(clienti_staff, use_container_width=True, hide_index=True)

        section("🛠️","Ultimi lavori registrati")
        if lavori_staff.empty:
            st.info("Nessun lavoro registrato.")
        else:
            st.dataframe(lavori_staff.head(20), use_container_width=True, hide_index=True)

        section("💬","Ultimi feedback ricevuti")
        if feedback_staff.empty:
            st.info("Nessun feedback ricevuto.")
        else:
            st.dataframe(feedback_staff.head(20), use_container_width=True, hide_index=True)
        return

    clienti = read_df("SELECT * FROM clienti")
    contratti = read_df("SELECT * FROM contratti")
    p = payments_df()
    fatt = read_df("SELECT * FROM fatture")

    clienti_totali = len(clienti)
    clienti_con_contratti = int(contratti["cliente_id"].nunique()) if not contratti.empty and "cliente_id" in contratti.columns else 0
    contratti_totali = len(contratti)

    if contratti.empty:
        totale_contratti = iva_contratti = totale_iva_inclusa = incasso_medio = 0.0
    else:
        imp = pd.to_numeric(contratti.get("importo_totale", 0), errors="coerce").fillna(0)
        iva_pct = pd.to_numeric(contratti.get("iva_percentuale", 22), errors="coerce").fillna(22)
        durata = pd.to_numeric(contratti.get("durata_mesi", 12), errors="coerce").replace(0, 12).fillna(12)
        totale_contratti = float(imp.sum())
        iva_contratti = float((imp * iva_pct / 100).sum())
        totale_iva_inclusa = totale_contratti + iva_contratti
        incasso_medio = float(((imp * (1 + iva_pct / 100)) / durata).sum())

    residuo_rate = float(p[p['stato'].isin(['Da pagare','Acconto'])]['residuo'].sum()) if not p.empty else 0.0
    scad_90 = len(contratti[(contratti['data_scadenza'] <= (date.today()+timedelta(days=90)).isoformat()) & (contratti['data_scadenza'] >= today_iso())]) if not contratti.empty else 0

    # Prima riga: conteggi separati, così non si confondono clienti e contratti.
    a,b,c,d = st.columns(4)
    a.metric("Clienti totali CRM", clienti_totali)
    b.metric("Clienti con contratti", clienti_con_contratti)
    c.metric("Contratti totali", contratti_totali)
    d.metric("Fatture", len(fatt))

    # Seconda riga: valori economici.
    c1,c2,c3,c4,c5 = st.columns(5)
    c1.metric("Totale contratti", money(totale_contratti))
    c2.metric("IVA contratti", money(iva_contratti))
    c3.metric("Totale contratti IVA inclusa", money(totale_iva_inclusa))
    c4.metric("Incasso medio mensile", money(incasso_medio))
    c5.metric("Residuo rate", money(residuo_rate))

    c6, c7 = st.columns(2)
    c6.metric("Scadenze 90 gg", scad_90)
    c7.metric("Rate aperte", len(p[p['stato'].isin(['Da pagare','Acconto','Sollecitata'])]) if not p.empty else 0)

    st.divider(); section("📚","Ultimi contratti")
    df = read_df("SELECT c.id, cl.ragione_sociale cliente, COALESCE(s.nome || ' ' || COALESCE(s.cognome,''), 'Non assegnato') staff, c.titolo, c.data_decorrenza, c.data_scadenza, c.importo_totale, c.iva_percentuale, c.stato FROM contratti c JOIN clienti cl ON cl.id=c.cliente_id LEFT JOIN staff s ON s.id=c.staff_id ORDER BY c.id DESC LIMIT 10")
    if df.empty: st.info("Nessun contratto creato.")
    else:
        df["iva_contratto"] = pd.to_numeric(df["importo_totale"], errors="coerce").fillna(0) * pd.to_numeric(df["iva_percentuale"], errors="coerce").fillna(22) / 100
        df["totale_iva_inclusa"] = pd.to_numeric(df["importo_totale"], errors="coerce").fillna(0) + df["iva_contratto"]
        for col in ["importo_totale", "iva_contratto", "totale_iva_inclusa"]:
            df[col] = df[col].apply(money)
        st.dataframe(df, use_container_width=True, hide_index=True)
    section("💶","Prossime rate da incassare")
    if p.empty:
        st.success("Nessuna rata.")
    else:
        # Mostra tutte le rate aperte del periodo selezionato, anche se già scadute.
        # Prima filtravo solo data >= oggi: per questo maggio 2026 risultava vuoto anche con rate presenti.
        aperte = p[(p["stato"] != "Pagata") & (p["residuo"].fillna(0) > 0.01)].copy()
        aperte = aperte.sort_values(["data_scadenza", "cliente", "numero_rata"])
        if aperte.empty:
            st.success("Non ci sono rate aperte da incassare.")
        else:
            mesi_map = {
                "01": "Gennaio", "02": "Febbraio", "03": "Marzo", "04": "Aprile",
                "05": "Maggio", "06": "Giugno", "07": "Luglio", "08": "Agosto",
                "09": "Settembre", "10": "Ottobre", "11": "Novembre", "12": "Dicembre"
            }
            anni_presenti = sorted({str(x)[:4] for x in aperte["data_scadenza"].astype(str) if len(str(x)) >= 4})
            mesi_presenti = sorted({str(x)[5:7] for x in aperte["data_scadenza"].astype(str) if len(str(x)) >= 7})

            c_sel1, c_sel2 = st.columns(2)
            anno_sel = c_sel1.selectbox("Selettore anno", ["Tutti gli anni"] + anni_presenti, key="dash_prossime_rate_anno")
            opzioni_mese = ["Tutti i mesi"] + [f"{m} - {mesi_map.get(m, m)}" for m in mesi_presenti]
            mese_sel = c_sel2.selectbox("Selettore mese", opzioni_mese, key="dash_prossime_rate_mese")

            filtrate = aperte.copy()
            if anno_sel != "Tutti gli anni":
                filtrate = filtrate[filtrate["data_scadenza"].astype(str).str[:4] == anno_sel]
            if mese_sel != "Tutti i mesi":
                mese_val = mese_sel.split(" - ")[0]
                filtrate = filtrate[filtrate["data_scadenza"].astype(str).str[5:7] == mese_val]

            if filtrate.empty:
                st.info("Non ci sono rate aperte nel periodo selezionato.")
            else:
                v = filtrate[["cliente","contratto","numero_rata","data_scadenza","imponibile","iva","totale","pagato","residuo","stato"]].copy().head(50)
                v = v.rename(columns={
                    "cliente": "Cliente",
                    "contratto": "Contratto",
                    "numero_rata": "Rata",
                    "data_scadenza": "Scadenza",
                    "imponibile": "Importo netto",
                    "iva": "IVA rata",
                    "totale": "Totale IVA inclusa",
                    "pagato": "Pagato",
                    "residuo": "Residuo",
                    "stato": "Stato"
                })
                for col in ["Importo netto","IVA rata","Totale IVA inclusa","Pagato","Residuo"]:
                    v[col] = v[col].apply(money)
                st.dataframe(v, use_container_width=True, hide_index=True)

    with st.expander("Rate scadute / da verificare", expanded=False):
        if p.empty:
            st.info("Nessuna rata.")
        else:
            overdue = p[(p["stato"] != "Pagata") & (p["residuo"].fillna(0) > 0.01) & (p["data_scadenza"].astype(str) < date.today().isoformat())].copy()
            overdue = overdue.sort_values(["data_scadenza", "cliente", "numero_rata"])
            if overdue.empty:
                st.success("Nessuna rata scaduta.")
            else:
                ov = overdue[["cliente","contratto","numero_rata","data_scadenza","imponibile","iva","totale","pagato","residuo","stato"]].copy()
                ov = ov.rename(columns={
                    "cliente": "Cliente",
                    "contratto": "Contratto",
                    "numero_rata": "Rata",
                    "data_scadenza": "Scadenza",
                    "imponibile": "Importo netto",
                    "iva": "IVA rata",
                    "totale": "Totale IVA inclusa",
                    "pagato": "Pagato",
                    "residuo": "Residuo",
                    "stato": "Stato"
                })
                for col in ["Importo netto","IVA rata","Totale IVA inclusa","Pagato","Residuo"]:
                    ov[col] = ov[col].apply(money)
                st.dataframe(ov, use_container_width=True, hide_index=True)

def page_aziende():
    header(); section("🏢","Aziende / emittenti","Modifica nomi aziende, dati fiscali, PEC, SDI e logo. Il sistema è rifattibile.")
    df = read_df("SELECT * FROM aziende ORDER BY is_default DESC, nome")
    st.dataframe(df[["id","nome","forma_giuridica","piva","cf","pec","codice_sdi","is_default"]], use_container_width=True, hide_index=True)
    with st.expander("➕ Crea / modifica azienda", expanded=True):
        ids = [0] + df["id"].astype(int).tolist()
        choice = st.selectbox("Azienda da modificare oppure 0 per nuova", ids, format_func=lambda x: "Nuova azienda" if x==0 else f"ID {x} - {df[df.id==x].iloc[0]['nome']}")
        row = {} if choice==0 else df[df.id==choice].iloc[0].to_dict()
        with st.form("azienda_form"):
            c1,c2 = st.columns(2)
            with c1:
                nome = st.text_input("Nome azienda *", row.get("nome", "")); forma = st.text_input("Forma giuridica", row.get("forma_giuridica", "")); piva = st.text_input("P.IVA", row.get("piva", "")); cf = st.text_input("Codice fiscale", row.get("cf", "")); pec = st.text_input("PEC", row.get("pec", "")); sdi = st.text_input("Codice SDI", row.get("codice_sdi", ""))
            with c2:
                sede = st.text_area("Sede", row.get("sede", "")); iban = st.text_input("IBAN", row.get("iban", "")); tel = st.text_input("Telefono", row.get("telefono", "")); email = st.text_input("Email", row.get("email", "")); is_def = st.checkbox("Azienda predefinita", bool(row.get("is_default", 0))) ; note = st.text_area("Note", row.get("note", ""))
            up = st.file_uploader("Logo azienda (opzionale)", type=["png","jpg","jpeg"])
            submit = st.form_submit_button("Salva azienda")
            if submit:
                logo = save_upload(up, DOC_DIR, "logo_") if up else row.get("logo_file")
                if is_def: execute("UPDATE aziende SET is_default=0")
                if choice==0:
                    execute("INSERT INTO aziende (nome,forma_giuridica,piva,cf,sede,pec,codice_sdi,iban,telefono,email,logo_file,note,is_default,created_at) VALUES (?,?,?,?,?,?,?,?,?,?,?,?,?,?)", (nome,forma,piva,cf,sede,pec,sdi,iban,tel,email,logo,note,1 if is_def else 0,now_iso()))
                else:
                    execute("UPDATE aziende SET nome=?,forma_giuridica=?,piva=?,cf=?,sede=?,pec=?,codice_sdi=?,iban=?,telefono=?,email=?,logo_file=?,note=?,is_default=? WHERE id=?", (nome,forma,piva,cf,sede,pec,sdi,iban,tel,email,logo,note,1 if is_def else 0,choice))
                st.success("Azienda salvata."); st.rerun()

def page_clienti():
    header(); section("👥","Clienti CRM","Scheda cliente completa, modificabile e con tab documenti/lavori/contratti.")
    if user_is_admin():
        df = read_df("SELECT * FROM clienti ORDER BY ragione_sociale")
        tab1, tab2 = st.tabs(["CRM cliente", "Nuovo / modifica cliente"])
    else:
        assigned = assigned_clients_df()
        if assigned.empty:
            st.info("Non hai ancora clienti assegnati.")
            return
        ids_visibili = ",".join(str(int(x)) for x in assigned["id"].tolist())
        df = read_df(f"SELECT * FROM clienti WHERE id IN ({ids_visibili}) ORDER BY ragione_sociale")
        tab1 = st.container()
        tab2 = None
    if user_is_admin() and tab2 is not None:
        with tab2:
            ids = [0] + df["id"].astype(int).tolist() if not df.empty else [0]
            choice = st.selectbox("Cliente da modificare oppure 0 per nuovo", ids, format_func=lambda x: "Nuovo cliente" if x==0 else f"ID {x} - {df[df.id==x].iloc[0]['ragione_sociale']}")
            row = {} if choice==0 else df[df.id==choice].iloc[0].to_dict()
            with st.form("cliente_form"):
                c1,c2,c3 = st.columns(3)
                with c1:
                    rag = st.text_input("Ragione sociale *", row.get("ragione_sociale", "")); forma = st.text_input("Forma giuridica", row.get("forma_giuridica", "")); piva = st.text_input("Partita IVA", row.get("partita_iva", "")); cf = st.text_input("Codice fiscale", row.get("codice_fiscale", "")); rea = st.text_input("REA", row.get("rea", "")); ateco = st.text_input("Codice Ateco", row.get("codice_ateco", ""))
                with c2:
                    sede = st.text_area("Sede legale", row.get("sede_legale", "")); pec = st.text_input("PEC", row.get("pec", "")); sdi = st.text_input("Codice SDI", row.get("codice_sdi", "")); leg = st.text_input("Titolare / legale rappresentante", row.get("legale_rappresentante", ""))
                with c3:
                    tel = st.text_input("Telefono", row.get("telefono", "")); email = st.text_input("Email", row.get("email", "")); settore = st.text_input("Settore", row.get("settore", "")); stato = st.selectbox("Stato CRM", ["Attivo","In trattativa","Sospeso","Ex cliente"], index=["Attivo","In trattativa","Sospeso","Ex cliente"].index(row.get("stato_crm","Attivo") if row.get("stato_crm","Attivo") in ["Attivo","In trattativa","Sospeso","Ex cliente"] else "Attivo")); note = st.text_area("Note", row.get("note", ""))
                if st.form_submit_button("Salva cliente"):
                    if choice==0: execute("INSERT INTO clienti (ragione_sociale,forma_giuridica,partita_iva,codice_fiscale,rea,sede_legale,pec,codice_sdi,legale_rappresentante,telefono,email,codice_ateco,settore,stato_crm,note,created_at) VALUES (?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?)", (rag,forma,piva,cf,rea,sede,pec,sdi,leg,tel,email,ateco,settore,stato,note,now_iso()))
                    else: execute("UPDATE clienti SET ragione_sociale=?,forma_giuridica=?,partita_iva=?,codice_fiscale=?,rea=?,sede_legale=?,pec=?,codice_sdi=?,legale_rappresentante=?,telefono=?,email=?,codice_ateco=?,settore=?,stato_crm=?,note=? WHERE id=?", (rag,forma,piva,cf,rea,sede,pec,sdi,leg,tel,email,ateco,settore,stato,note,choice))
                    st.success("Cliente salvato."); st.rerun()
    with tab1:
        if df.empty: st.info("Inserisci prima un cliente."); return
        label, opts = client_select("Seleziona cliente CRM", "crm_cliente")
        cid = opts[label]
        row = read_df("SELECT * FROM clienti WHERE id=?", (cid,)).iloc[0].to_dict()
        c1,c2,c3,c4 = st.columns(4)
        with c1: card("👤","Cliente", row.get("ragione_sociale"))
        with c2: card("🧾","P.IVA", row.get("partita_iva") or "-")
        with c3: card("📨","PEC", row.get("pec") or "-")
        with c4: card("📌","Stato", row.get("stato_crm") or "-")
        if user_is_admin() or staff_access_level() == "Gestione Finanziaria":
            t1,t2,t4,t5,t6 = st.tabs(["📋 Dati", "📚 Contratti", "📎 Documenti", "🧾 Fatture", "💬 Feedback"])
            t3 = None
        else:
            t1,t3,t4,t6 = st.tabs(["📋 Dati", "🛠️ Lavori", "📎 Documenti", "💬 Feedback"])
            t2 = t5 = None
        with t1:
            st.markdown("### Modifica dati cliente")
            with st.form(f"crm_edit_cliente_{cid}"):
                ec1, ec2, ec3 = st.columns(3)
                with ec1:
                    rag = st.text_input("Ragione sociale *", row.get("ragione_sociale", ""), key=f"crm_rag_{cid}")
                    forma = st.text_input("Forma giuridica", row.get("forma_giuridica", ""), key=f"crm_forma_{cid}")
                    piva = st.text_input("Partita IVA", row.get("partita_iva", ""), key=f"crm_piva_{cid}")
                    cf = st.text_input("Codice fiscale", row.get("codice_fiscale", ""), key=f"crm_cf_{cid}")
                    rea = st.text_input("REA", row.get("rea", "") or "", key=f"crm_rea_{cid}")
                with ec2:
                    sede = st.text_area("Sede legale", row.get("sede_legale", "") or "", key=f"crm_sede_{cid}")
                    pec = st.text_input("PEC", row.get("pec", "") or "", key=f"crm_pec_{cid}")
                    sdi = st.text_input("Codice SDI", row.get("codice_sdi", "") or "", key=f"crm_sdi_{cid}")
                    ateco = st.text_input("Codice Ateco", row.get("codice_ateco", "") or "", key=f"crm_ateco_{cid}")
                    leg = st.text_input("Titolare / Legale rappresentante", row.get("legale_rappresentante", "") or "", key=f"crm_leg_{cid}")
                with ec3:
                    tel = st.text_input("Telefono", row.get("telefono", "") or "", key=f"crm_tel_{cid}")
                    email = st.text_input("Email", row.get("email", "") or "", key=f"crm_email_{cid}")
                    settore = st.text_input("Settore", row.get("settore", "") or "", key=f"crm_settore_{cid}")
                    stato_attuale = row.get("stato_crm", "Attivo") or "Attivo"
                    stati = ["Attivo","In trattativa","Sospeso","Ex cliente"]
                    stato_idx = stati.index(stato_attuale) if stato_attuale in stati else 0
                    stato = st.selectbox("Stato CRM", stati, index=stato_idx, key=f"crm_stato_{cid}")
                    note = st.text_area("Note", row.get("note", "") or "", key=f"crm_note_{cid}")
                if st.form_submit_button("💾 Salva modifiche cliente"):
                    if not str(rag).strip():
                        st.error("La ragione sociale è obbligatoria.")
                    else:
                        execute(
                            "UPDATE clienti SET ragione_sociale=?,forma_giuridica=?,partita_iva=?,codice_fiscale=?,rea=?,sede_legale=?,pec=?,codice_sdi=?,legale_rappresentante=?,telefono=?,email=?,codice_ateco=?,settore=?,stato_crm=?,note=? WHERE id=?",
                            (rag,forma,piva,cf,rea,sede,pec,sdi,leg,tel,email,ateco,settore,stato,note,cid)
                        )
                        st.success("Dati cliente aggiornati.")
                        st.rerun()
        if (user_is_admin() or staff_access_level() == "Gestione Finanziaria") and t2 is not None:
            with t2:
                contracts = read_df("SELECT id,titolo,data_decorrenza,data_scadenza,importo_totale,stato,file_pdf,file_docx FROM contratti WHERE cliente_id=? ORDER BY id DESC", (cid,))
                if contracts.empty: st.info("Nessun contratto.")
                else:
                    v=contracts.copy(); v["importo_totale"]=v["importo_totale"].apply(money); st.dataframe(v, use_container_width=True, hide_index=True)
        if t3 is not None:
            with t3:
                lavori_cliente(cid)
        with t4:
            documenti_cliente(cid)
        if (user_is_admin() or staff_access_level() == "Gestione Finanziaria") and t5 is not None:
            with t5:
                inv = read_df("SELECT id,numero,data_fattura,descrizione,totale,stato,file_pdf FROM fatture WHERE cliente_id=? ORDER BY id DESC", (cid,))
                if inv.empty: st.info("Nessuna fattura.")
                else:
                    inv["totale"] = inv["totale"].apply(money); st.dataframe(inv, use_container_width=True, hide_index=True)


def feedback_cliente(cid: int, contratto_id: int | None = None):
    user = current_staff_user()
    with st.form(f"feedback_form_{cid}_{contratto_id}"):
        c1,c2,c3 = st.columns(3)
        data_fb = c1.date_input("Data feedback", value=date.today(), key=f"fb_data_{cid}_{contratto_id}")
        provenienza = c2.selectbox("Provenienza", ["Cliente","Interno","WhatsApp","Email","Telefonata","Riunione"], key=f"fb_prov_{cid}_{contratto_id}")
        valutazione = c3.selectbox("Valutazione", ["Non indicata",1,2,3,4,5], key=f"fb_val_{cid}_{contratto_id}")
        testo = st.text_area("Testo feedback")
        note = st.text_area("Note interne feedback")
        up = st.file_uploader("Allega documento/foto feedback", type=["pdf","png","jpg","jpeg","docx","txt"], key=f"fb_up_{cid}_{contratto_id}")
        if st.form_submit_button("💬 Salva feedback"):
            rel = save_upload(up, DOC_DIR, f"feedback_cliente_{cid}_") if up else None
            val = None if valutazione == "Non indicata" else int(valutazione)
            execute("INSERT INTO feedback_clienti (cliente_id,contratto_id,staff_id,data_feedback,provenienza,valutazione,testo_feedback,allegato_file,note,created_at) VALUES (?,?,?,?,?,?,?,?,?,?)",
                    (cid, contratto_id, int(user["id"]) if user else None, data_fb.isoformat(), provenienza, val, testo, rel, note, now_iso()))
            st.success("Feedback salvato.")
            st.rerun()

    if user_is_admin() or staff_access_level() == "Gestione Finanziaria":
        fb = read_df("SELECT f.id,f.data_feedback,f.provenienza,f.valutazione,f.testo_feedback,COALESCE(s.nome || ' ' || COALESCE(s.cognome,''),'Non indicato') staff,f.allegato_file,f.note FROM feedback_clienti f LEFT JOIN staff s ON s.id=f.staff_id WHERE f.cliente_id=? AND (? IS NULL OR f.contratto_id=?) ORDER BY f.data_feedback DESC,id DESC", (cid, contratto_id, contratto_id))
    else:
        fb = read_df("SELECT f.id,f.data_feedback,f.provenienza,f.valutazione,f.testo_feedback,COALESCE(s.nome || ' ' || COALESCE(s.cognome,''),'Non indicato') staff,f.allegato_file,f.note FROM feedback_clienti f LEFT JOIN staff s ON s.id=f.staff_id JOIN contratti c ON c.cliente_id=f.cliente_id WHERE f.cliente_id=? AND c.staff_id=? AND (? IS NULL OR f.contratto_id=?) GROUP BY f.id ORDER BY f.data_feedback DESC,id DESC", (cid, int(user["id"]) if user else 0, contratto_id, contratto_id))
    if fb.empty:
        st.info("Nessun feedback registrato.")
    else:
        st.dataframe(fb, use_container_width=True, hide_index=True)


        with t6:
            feedback_cliente(cid)

def documenti_cliente(cid: int, contratto_id: int | None = None):
    with st.form(f"doc_form_{cid}_{contratto_id}"):
        c1,c2 = st.columns(2)
        titolo = c1.text_input("Titolo documento")
        tipo = c2.selectbox("Tipo", ["Contratto","Documento cliente","Pagamento","Lavoro","Fattura","Altro"])
        note = st.text_area("Note documento")
        up = st.file_uploader("Allega PDF/foto/documento", type=["pdf","png","jpg","jpeg","docx","xlsx","csv"])
        if st.form_submit_button("Carica documento"):
            rel = save_upload(up, DOC_DIR, f"cliente_{cid}_") if up else None
            execute("INSERT INTO documenti (cliente_id,contratto_id,tipo,titolo,file_path,note,created_at) VALUES (?,?,?,?,?,?,?)", (cid,contratto_id,tipo,titolo,rel,note,now_iso()))
            st.success("Documento caricato."); st.rerun()
    docs = read_df("SELECT id,tipo,titolo,file_path,note,created_at FROM documenti WHERE cliente_id=? AND (? IS NULL OR contratto_id=?) ORDER BY id DESC", (cid, contratto_id, contratto_id))
    if docs.empty: st.info("Nessun documento.")
    else: st.dataframe(docs, use_container_width=True, hide_index=True)


def sintesi_tre_righe_testo(testo: str) -> str:
    """Restituisce una sintesi leggibile in massimo 3 frasi/righe senza usare AI."""
    testo = re.sub(r"\s+", " ", str(testo or "")).strip()
    if not testo:
        return "Attività registrata nel gestionale cliente e collegata al periodo selezionato.\nIl lavoro è stato tracciato come avanzamento operativo del progetto.\nLa scheda rimane disponibile nel CRM con eventuali allegati e aggiornamenti successivi."
    # Spezza prima per punteggiatura, poi ricompone al massimo 3 segmenti.
    parti = [p.strip() for p in re.split(r"(?<=[\.\!\?;])\s+", testo) if p.strip()]
    if len(parti) < 3:
        parole = testo.split()
        if len(parole) > 42:
            parti = [" ".join(parole[:18]), " ".join(parole[18:36]), " ".join(parole[36:54])]
        else:
            parti = [testo]
    return "\n".join(parti[:3])


def genera_report_lavori_pdf(cliente: dict, lavori: pd.DataFrame, data_da: date, data_a: date, includi_note: bool = False, includi_allegati: bool = True) -> Path:
    """Genera un report PDF dei lavori filtrati, con descrizione sintetica di 3 righe per lavoro."""
    safe_name = slug(str(cliente.get("ragione_sociale") or "cliente"))
    out = REPORT_DIR / f"report_lavori_{safe_name}_{data_da.isoformat()}_{data_a.isoformat()}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.pdf"

    doc = SimpleDocTemplate(str(out), pagesize=A4, rightMargin=1.6*cm, leftMargin=1.6*cm, topMargin=1.5*cm, bottomMargin=1.5*cm)
    styles = getSampleStyleSheet()
    title_style = ParagraphStyle("ReportTitle", parent=styles["Title"], fontSize=18, leading=22, alignment=TA_LEFT, spaceAfter=10)
    h_style = ParagraphStyle("ReportH", parent=styles["Heading2"], fontSize=12, leading=15, textColor=colors.HexColor("#0c1d2f"), spaceBefore=8, spaceAfter=6)
    body = ParagraphStyle("ReportBody", parent=styles["BodyText"], fontSize=9.3, leading=12, spaceAfter=6)
    small = ParagraphStyle("ReportSmall", parent=styles["BodyText"], fontSize=8.3, leading=10, textColor=colors.HexColor("#53657a"))

    elems = []
    elems.append(Paragraph("Report lavori svolti", title_style))
    elems.append(Paragraph(f"<b>Cliente:</b> {html.escape(str(cliente.get('ragione_sociale') or ''))}", body))
    elems.append(Paragraph(f"<b>Periodo:</b> {data_da.strftime('%d/%m/%Y')} - {data_a.strftime('%d/%m/%Y')}", body))
    elems.append(Paragraph(f"<b>Totale lavori nel periodo:</b> {len(lavori)}", body))
    elems.append(Spacer(1, 0.15*cm))
    elems.append(Paragraph(
        "Nel periodo selezionato sono state riepilogate le attività operative registrate nel CRM. "
        "Il report documenta i lavori svolti, lo stato di avanzamento e gli eventuali allegati collegati, "
        "così da poter condividere con il cliente una sintesi chiara e professionale del lavoro effettuato.",
        body
    ))
    elems.append(Spacer(1, 0.25*cm))

    if lavori.empty:
        elems.append(Paragraph("Nessun lavoro registrato nel periodo selezionato.", body))
    else:
        for _, r in lavori.iterrows():
            titolo = html.escape(str(r.get("titolo") or "Lavoro senza titolo"))
            tipo = html.escape(str(r.get("tipo_lavoro") or "-"))
            stato = html.escape(str(r.get("stato") or "-"))
            data_lav = html.escape(str(r.get("data_lavoro") or "-"))
            desc = html.escape(sintesi_tre_righe_testo(str(r.get("descrizione") or ""))).replace("\n", "<br/>")

            elems.append(Paragraph(f"{data_lav} · {tipo} · {titolo}", h_style))
            elems.append(Paragraph(f"<b>Stato:</b> {stato}", small))
            elems.append(Paragraph(desc, body))

            if includi_note and str(r.get("note_interne") or "").strip():
                note = html.escape(sintesi_tre_righe_testo(str(r.get("note_interne") or ""))).replace("\n", "<br/>")
                elems.append(Paragraph(f"<b>Note interne:</b><br/>{note}", small))

            if includi_allegati and str(r.get("allegato_file") or "").strip() and str(r.get("allegato_file")).lower() != "none":
                elems.append(Paragraph(f"<b>Allegato:</b> {html.escape(str(r.get('allegato_file')))}", small))

            elems.append(Spacer(1, 0.12*cm))

    doc.build(elems)
    return out


def lavori_cliente(cid: int, contratto_id: int | None = None):
    cliente_df = read_df("SELECT * FROM clienti WHERE id=?", (cid,))
    cliente = cliente_df.iloc[0].to_dict() if not cliente_df.empty else {"ragione_sociale": "Cliente"}

    with st.form(f"lav_form_{cid}_{contratto_id}"):
        c1,c2,c3 = st.columns(3)
        data_lav = c1.date_input("Data lavoro", value=date.today())
        tipo = c2.selectbox("Tipo lavoro", ["Sito web","Social","ADS","Grafica","Marketplace","Consulenza","Documenti","Contratto","Altro"])
        stato = c3.selectbox("Stato", ["Da fare","In lavorazione","Completato","Consegnato","In attesa cliente","Bloccato"])
        titolo = st.text_input("Titolo lavoro")
        desc = st.text_area("Descrizione lavoro")
        note = st.text_area("Note interne")
        up = st.file_uploader("Allega PDF/foto lavoro", type=["pdf","png","jpg","jpeg","docx","xlsx","csv"])
        if st.form_submit_button("Salva lavoro"):
            rel = save_upload(up, WORK_DIR, f"lavoro_cliente_{cid}_") if up else None
            execute("INSERT INTO lavori (cliente_id,contratto_id,data_lavoro,tipo_lavoro,titolo,descrizione,stato,allegato_file,note_interne,created_at) VALUES (?,?,?,?,?,?,?,?,?,?)", (cid,contratto_id,data_lav.isoformat(),tipo,titolo,desc,stato,rel,note,now_iso()))
            if rel: execute("INSERT INTO documenti (cliente_id,contratto_id,tipo,titolo,file_path,note,created_at) VALUES (?,?,?,?,?,?,?)", (cid,contratto_id,"Lavoro",titolo,rel,"Allegato lavoro",now_iso()))
            st.success("Lavoro salvato."); st.rerun()

    st.markdown("### Filtra lavori per periodo")
    base = read_df(
        "SELECT id,data_lavoro,tipo_lavoro,titolo,descrizione,stato,allegato_file,note_interne FROM lavori WHERE cliente_id=? AND (? IS NULL OR contratto_id=?) ORDER BY data_lavoro DESC,id DESC",
        (cid, contratto_id, contratto_id)
    )

    if base.empty:
        data_min = date.today().replace(day=1)
        data_max = date.today()
    else:
        dates = pd.to_datetime(base["data_lavoro"], errors="coerce").dropna()
        data_min = dates.min().date() if not dates.empty else date.today().replace(day=1)
        data_max = dates.max().date() if not dates.empty else date.today()

    f1,f2,f3,f4 = st.columns([1,1,1,1])
    data_da = f1.date_input("Dal", value=data_min, key=f"lav_da_{cid}_{contratto_id}")
    data_a = f2.date_input("Al", value=data_max, key=f"lav_a_{cid}_{contratto_id}")
    tipo_filter = f3.selectbox("Tipo", ["Tutti","Sito web","Social","ADS","Grafica","Marketplace","Consulenza","Documenti","Contratto","Altro"], key=f"lav_tipo_{cid}_{contratto_id}")
    stato_filter = f4.selectbox("Stato", ["Tutti","Da fare","In lavorazione","Completato","Consegnato","In attesa cliente","Bloccato"], key=f"lav_stato_{cid}_{contratto_id}")

    lavori = base.copy()
    if not lavori.empty:
        lavori["_data"] = pd.to_datetime(lavori["data_lavoro"], errors="coerce").dt.date
        lavori = lavori[(lavori["_data"] >= data_da) & (lavori["_data"] <= data_a)]
        if tipo_filter != "Tutti":
            lavori = lavori[lavori["tipo_lavoro"].astype(str) == tipo_filter]
        if stato_filter != "Tutti":
            lavori = lavori[lavori["stato"].astype(str) == stato_filter]
        lavori = lavori.drop(columns=["_data"], errors="ignore")

    csum1, csum2, csum3, csum4 = st.columns(4)
    with csum1: card("🛠️", "Lavori nel periodo", str(len(lavori)))
    with csum2: card("✅", "Completati", str(len(lavori[lavori["stato"].astype(str)=="Completato"])) if not lavori.empty else "0")
    with csum3: card("📎", "Con allegato", str(len(lavori[lavori["allegato_file"].notna() & (lavori["allegato_file"].astype(str)!="")])) if not lavori.empty else "0")
    with csum4: card("📅", "Intervallo", f"{data_da.strftime('%d/%m/%Y')} - {data_a.strftime('%d/%m/%Y')}")

    st.markdown("### Report lavori")
    r1, r2, r3 = st.columns([1,1,2])
    includi_note = r1.checkbox("Includi note interne", value=False, key=f"rep_note_{cid}_{contratto_id}")
    includi_allegati = r2.checkbox("Includi riferimenti allegati", value=True, key=f"rep_all_{cid}_{contratto_id}")
    if r3.button("🧾 Genera report lavori PDF", key=f"gen_report_lav_{cid}_{contratto_id}"):
        if lavori.empty:
            st.warning("Non ci sono lavori nel periodo selezionato.")
        else:
            pdf_path = genera_report_lavori_pdf(cliente, lavori, data_da, data_a, includi_note, includi_allegati)
            st.session_state[f"ultimo_report_lavori_{cid}_{contratto_id}"] = str(pdf_path)
            st.success("Report lavori generato.")

    report_key = f"ultimo_report_lavori_{cid}_{contratto_id}"
    if st.session_state.get(report_key):
        pdf_path = Path(st.session_state[report_key])
        if pdf_path.exists():
            st.download_button("⬇️ Scarica report lavori PDF", data=pdf_path.read_bytes(), file_name=pdf_path.name, mime="application/pdf", key=f"down_report_lav_{cid}_{contratto_id}")

    st.markdown("### Storico lavori")
    if lavori.empty:
        st.info("Nessun lavoro registrato nel periodo selezionato.")
    else:
        st.dataframe(
            lavori[["id","data_lavoro","tipo_lavoro","titolo","descrizione","stato","allegato_file","note_interne"]],
            use_container_width=True,
            hide_index=True
        )


def page_importa_contratto():
    header(); section("📥", "Importa contratto", "Carica PDF o foto da qualunque punto del gestionale e trasferisci i dati nella creazione contratto.")
    st.info("Questa voce resta sempre disponibile nel menu laterale. Dopo l'importazione potrai completare, modificare e generare il contratto dalla sezione Crea nuovo contratto.")
    blocco_importa_contratto("pagina_importa")

def page_crea_contratto():
    header(); section("📝","Crea nuovo contratto","Campi manuali, importazione PDF/foto, servizi liberi e base contrattuale modificabile.")
    aziende = read_df("SELECT * FROM aziende")
    clienti = read_df("SELECT * FROM clienti")
    tabs = st.tabs(["Importa PDF/foto", "Contratto", "Servizi", "Anteprima / creazione"])

    if "new_services" not in st.session_state:
        st.session_state.new_services = pd.DataFrame([{"Servizio":"", "Descrizione":""}])
    if "import_contract_data" not in st.session_state:
        st.session_state.import_contract_data = {}

    def applica_import_nei_campi(dati: dict):
        """Copia i dati letti dal PDF/foto nei campi manuali della scheda Contratto."""
        st.session_state["new_cliente_mode"] = "Nuovo cliente manuale"
        st.session_state["new_rag"] = dati.get("ragione_sociale", "")
        st.session_state["new_piva"] = dati.get("partita_iva", "")
        st.session_state["new_cf"] = dati.get("codice_fiscale", "")
        st.session_state["new_sede"] = dati.get("sede_legale", "")
        st.session_state["new_pec"] = dati.get("pec", "")
        st.session_state["new_email"] = dati.get("email", "")
        st.session_state["new_tel"] = dati.get("telefono", "")
        st.session_state["new_leg"] = dati.get("legale_rappresentante", "")
        st.session_state["new_titolo"] = dati.get("titolo", "CONTRATTO DI CONSULENZA STRATEGICA E OPERATIVA")
        st.session_state["new_sottotitolo"] = dati.get("sottotitolo", "sviluppo digitale, visibilità estera, apertura mercati e richieste di appuntamento")
        st.session_state["new_tipo"] = dati.get("tipo_contratto", "Consulenza strategica e operativa")
        st.session_state["new_data_firma"] = parse_date(dati.get("data_firma"), date.today())
        st.session_state["new_luogo"] = dati.get("luogo_firma", "Napoli")
        st.session_state["new_decorrenza"] = parse_date(dati.get("data_decorrenza"), date.today())
        st.session_state["new_durata"] = int(dati.get("durata_mesi", 12) or 12)
        st.session_state["new_importo"] = float(dati.get("importo_totale", 6000.0) or 0.0)
        st.session_state["new_iva_pct"] = float(dati.get("iva_percentuale", 22.0) or 22.0)
        st.session_state["new_modalita"] = dati.get("modalita_pagamento", "Trimestrale")
        st.session_state["new_foro"] = dati.get("foro_competente", "Napoli")
        st.session_state["new_note"] = "Creato da importazione contratto."
        if dati.get("servizi"):
            st.session_state.new_services = pd.DataFrame(dati.get("servizi"))
        st.session_state["contratto_importato_nei_campi"] = True

    imp = st.session_state.import_contract_data or {}

    with tabs[0]:
        section("📥","Importa PDF/foto contratto","Carica un contratto già esistente: il sistema prova a leggere dati cliente, importo, data e servizi.")
        uploaded_contract = st.file_uploader("Carica PDF o foto contratto", type=["pdf","png","jpg","jpeg","webp"], key="upload_contract_import_new")
        cimp1, cimp2 = st.columns([1,2])
        with cimp1:
            leggi = st.button("📖 Leggi PDF/foto ed estrai i dati", use_container_width=True)
        with cimp2:
            st.info("Dopo la lettura, controlla sempre i campi nella scheda 'Contratto' prima di creare il documento.")
        if leggi:
            if uploaded_contract is None:
                st.error("Carica prima un PDF o una foto del contratto.")
            else:
                try:
                    raw = estrai_testo_da_upload_contratto(uploaded_contract)
                    dati = estrai_dati_contratto_da_testo(raw)
                    st.session_state.import_contract_data = dati
                    st.session_state.import_contract_file_name = uploaded_contract.name
                    st.session_state.import_contract_file_bytes = get_upload_bytes(uploaded_contract)
                    if dati.get("servizi"):
                        st.session_state.new_services = pd.DataFrame(dati.get("servizi"))
                    st.success("Dati estratti. Ora apri la scheda 'Contratto' e correggi/verifica prima di creare.")
                    sc1, sc2, sc3, sc4 = st.columns(4)
                    with sc1:
                        card("👤", "Cliente rilevato", dati.get("ragione_sociale") or "Da verificare")
                    with sc2:
                        card("🧾", "P.IVA / C.F.", dati.get("partita_iva") or dati.get("codice_fiscale") or "Da verificare")
                    with sc3:
                        card("📍", "Sede", dati.get("sede_legale") or "Da verificare")
                    with sc4:
                        card("💶", "Importo", money(dati.get("importo_totale")) if dati.get("importo_totale") else "Da verificare")
                    if dati.get("servizi"):
                        st.markdown("**Servizi rilevati dal contratto:**")
                        st.dataframe(pd.DataFrame(dati.get("servizi")), use_container_width=True, hide_index=True)
                    with st.expander("Anteprima testo letto", expanded=False):
                        st.text_area("Testo estratto", clean_import_preview_text(raw)[:15000], height=260)
                except Exception as exc:
                    st.error(str(exc))

        if st.session_state.get("import_contract_data"):
            st.divider()
            st.markdown("### ✅ Importa contratto nei campi")
            st.write("Dopo aver letto il file, clicca qui per trasferire i dati riconosciuti nella scheda **Contratto**. Potrai modificarli prima della creazione definitiva.")
            col_importa, col_pulisci = st.columns([1, 1])
            with col_importa:
                if st.button("✅ Importa contratto nei campi compilabili", use_container_width=True, key="btn_importa_contratto_nei_campi"):
                    applica_import_nei_campi(st.session_state.import_contract_data)
                    st.success("Contratto importato nei campi. Apri la scheda 'Contratto' per verificare e correggere i dati.")
                    st.rerun()
            with col_pulisci:
                if st.button("🧹 Pulisci importazione", use_container_width=True, key="btn_pulisci_import_contratto"):
                    for k in ["import_contract_data", "import_contract_file_name", "import_contract_file_bytes", "contratto_importato_nei_campi"]:
                        if k in st.session_state:
                            del st.session_state[k]
                    st.session_state.new_services = pd.DataFrame([{"Servizio":"", "Descrizione":""}])
                    st.rerun()

        if st.session_state.get("contratto_importato_nei_campi"):
            st.success("Contratto importato nei campi compilabili. Puoi andare nella scheda 'Contratto' e poi in 'Anteprima / creazione'.")

    # ricarica dopo eventuale import
    imp = st.session_state.import_contract_data or {}

    with tabs[1]:
        c1,c2 = st.columns(2)
        with c1:
            azi_label, azi_opts = company_select("Azienda emittente", "new_az")
            azienda_id = azi_opts[azi_label]
            templates_disponibili = read_df("SELECT id,nome,descrizione FROM templates_contratto WHERE COALESCE(attivo,1)=1 ORDER BY nome")
            if templates_disponibili.empty:
                st.warning("Nessun template attivo trovato: creane uno dalla sezione Template contratti.")
                template_id = None
            else:
                template_opts = {f"{r['nome']} · ID {r['id']}": int(r['id']) for _, r in templates_disponibili.iterrows()}
                template_label = st.selectbox("📄 Template da usare", list(template_opts.keys()), key="new_template_id_label")
                template_id = template_opts[template_label]
                descrizione_template = templates_disponibili[templates_disponibili['id'] == template_id].iloc[0].get('descrizione') or ""
                if descrizione_template:
                    st.caption(descrizione_template)
            if "new_cliente_mode" not in st.session_state:
                st.session_state["new_cliente_mode"] = "Cliente esistente"
            mode = st.radio("Cliente", ["Cliente esistente", "Nuovo cliente manuale"], horizontal=True, key="new_cliente_mode")
            if mode == "Cliente esistente" and not clienti.empty:
                cl_label, cl_opts = client_select("Seleziona cliente", "new_cl")
                cliente_id = cl_opts[cl_label]
                rag=piva=cf=sede=pec=leg=tel=email=""
            else:
                cliente_id = None
                st.markdown("**Nuovo cliente/contatto**")
                for key, default in {
                    "new_rag":"", "new_piva":"", "new_cf":"", "new_sede":"",
                    "new_pec":"", "new_email":"", "new_tel":"", "new_leg":""
                }.items():
                    st.session_state.setdefault(key, default)
                rag = st.text_input("Ragione sociale cliente *", key="new_rag")
                piva = st.text_input("P.IVA", key="new_piva")
                cf = st.text_input("Codice fiscale", key="new_cf")
                sede = st.text_area("Sede legale", key="new_sede")
                pec = st.text_input("PEC", key="new_pec")
                email = st.text_input("Email", key="new_email")
                tel = st.text_input("Telefono", key="new_tel")
                leg = st.text_input("Titolare / rappresentante", key="new_leg")
        with c2:
            st.session_state.setdefault("new_titolo", "CONTRATTO DI CONSULENZA STRATEGICA E OPERATIVA")
            st.session_state.setdefault("new_sottotitolo", "sviluppo digitale, visibilità estera, apertura mercati e richieste di appuntamento")
            st.session_state.setdefault("new_luogo", "Napoli")
            st.session_state.setdefault("new_data_firma", date.today())
            st.session_state.setdefault("new_decorrenza", date.today())
            st.session_state.setdefault("new_durata", 12)
            st.session_state.setdefault("new_importo", 6000.0)
            st.session_state.setdefault("new_iva_pct", 22.0)
            st.session_state.setdefault("new_foro", "Napoli")
            st.session_state.setdefault("new_note", "")
            titolo = st.text_input("Titolo contratto", key="new_titolo")
            sottotitolo = st.text_input("Sottotitolo", key="new_sottotitolo")
            tipi = ["Consulenza strategica e operativa","Gestione social","Sito web e mercati esteri","Marketplace","Joint venture operativa","Prestazione manager","Altro"]
            tipo_default = st.session_state.get("new_tipo", "Consulenza strategica e operativa")
            tipo = st.selectbox("Tipo contratto", tipi, index=tipi.index(tipo_default) if tipo_default in tipi else 0, key="new_tipo")
            data_firma = st.date_input("Data firma", key="new_data_firma")
            luogo = st.text_input("Luogo firma", key="new_luogo")
            decorrenza = st.date_input("Data decorrenza", key="new_decorrenza")
            durata = st.number_input("Durata mesi", min_value=1, max_value=60, step=1, key="new_durata")
            importo = st.number_input("Importo totale imponibile", min_value=0.0, step=100.0, key="new_importo")
            iva_pct = st.number_input("IVA %", min_value=0.0, max_value=100.0, step=1.0, key="new_iva_pct")
            pags = ["Mensile","Bimestrale","Trimestrale","Semestrale","Annuale"]
            modalita_default = st.session_state.get("new_modalita", "Trimestrale")
            modalita = st.selectbox("Modalità pagamento", pags, index=pags.index(modalita_default) if modalita_default in pags else 0, key="new_modalita")
            foro = st.text_input("Foro competente", key="new_foro")
            stato = st.selectbox("Stato", ["Bozza","Inviato","Firmato","Attivo","Sospeso","Scaduto","Archiviato"], key="new_stato")
            staff_label_new, staff_opts_new = staff_select("Responsabile staff", "new_staff_id", include_none=True)
            staff_id_new = staff_opts_new[staff_label_new]
            note = st.text_area("Note contratto", key="new_note")
    with tabs[2]:
        section("🧩","Servizi compresi","Compila a mano i servizi: saranno inseriti nell'ART. 2 del contratto.")
        st.session_state.new_services = st.data_editor(st.session_state.new_services, num_rows="dynamic", use_container_width=True, hide_index=True, key="serv_editor")
        extra = st.text_area("Clausole / note extra da inserire nel contratto")
    with tabs[3]:
        servizi = [r for r in st.session_state.new_services.to_dict("records") if str(r.get("Servizio","")).strip()]
        st.markdown("**Servizi che verranno generati nel contratto:**")
        st.write(servizi if servizi else "Nessun servizio inserito.")
        if st.button("➕ Crea nuovo contratto e genera Word/PDF"):
            if mode == "Nuovo cliente manuale":
                if not rag.strip(): st.error("Inserisci la ragione sociale del cliente."); return
                cliente_id = execute("INSERT INTO clienti (ragione_sociale,partita_iva,codice_fiscale,sede_legale,pec,email,telefono,legale_rappresentante,created_at) VALUES (?,?,?,?,?,?,?,?,?)", (rag,piva,cf,sede,pec,email,tel,leg,now_iso()))
            if not cliente_id: st.error("Seleziona o crea un cliente."); return
            scad = add_months(decorrenza, int(durata))
            if not template_id:
                st.error("Seleziona o crea un template contratto prima di generare il documento.")
                return
            cid = execute("""INSERT INTO contratti (azienda_id,cliente_id,template_id,staff_id,titolo,sottotitolo,tipo_contratto,data_firma,luogo_firma,data_decorrenza,data_scadenza,durata_mesi,importo_totale,iva_percentuale,modalita_pagamento,foro_competente,stato,servizi_json,clausole_extra,note,created_at) VALUES (?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?)""", (azienda_id,cliente_id,template_id,staff_id_new,titolo,sottotitolo,tipo,data_firma.isoformat(),luogo,decorrenza.isoformat(),scad.isoformat(),int(durata),float(importo),float(iva_pct),modalita,foro,stato,json.dumps(servizi, ensure_ascii=False),extra,note,now_iso()))
            # salva anche il PDF/foto importato come documento collegato al cliente/contratto
            if st.session_state.get("import_contract_file_bytes") and st.session_state.get("import_contract_file_name"):
                DOC_DIR.mkdir(parents=True, exist_ok=True)
                fname = f"{datetime.now().strftime('%Y%m%d_%H%M%S')}_contratto_importato_{slug(st.session_state.import_contract_file_name)}"
                fpath = DOC_DIR / fname
                fpath.write_bytes(st.session_state.import_contract_file_bytes)
                rel = str(fpath.relative_to(BASE_DIR))
                execute("INSERT INTO documenti (cliente_id,contratto_id,tipo,titolo,file_path,note,created_at) VALUES (?,?,?,?,?,?,?)", (cliente_id,cid,"Contratto importato","Contratto originale importato",rel,"Caricato dalla sezione Crea nuovo contratto",now_iso()))
            genera_rate(cid, decorrenza, int(durata), float(importo), float(iva_pct), modalita)
            docx, pdf = generate_contract_files(cid)
            st.success(f"Contratto creato e generato. ID: {cid}")
            file_download_button(docx, "Scarica Word"); file_download_button(pdf, "Scarica PDF")

def page_contratti():
    header(); section("📚","Contratti","Modifica contratti, servizi, tab pagamenti, lavori e documenti.")
    if not user_is_admin() and not can_see_contracts_operational():
        st.error("Accesso ai contratti non consentito per questo livello staff.")
        return
    if not user_is_admin() and can_see_contracts_operational():
        user = current_staff_user()
        df = read_df("""
            SELECT c.id, cl.ragione_sociale cliente, c.titolo, c.data_decorrenza, c.data_scadenza, c.stato
            FROM contratti c
            JOIN clienti cl ON cl.id=c.cliente_id
            WHERE c.staff_id=? OR ?=1
            ORDER BY c.id DESC
        """, (int(user["id"]), 1 if can_see_all_clients() else 0))
        if df.empty:
            st.info("Nessun contratto operativo visibile.")
            return
        st.dataframe(df, use_container_width=True, hide_index=True)
        opts = {f"ID {r['id']} · {r['cliente']} · {r['titolo']}": int(r['id']) for _,r in df.iterrows()}
        contratto_id = opts[st.selectbox("Seleziona contratto operativo", list(opts.keys()))]
        ctr = read_df("SELECT * FROM contratti WHERE id=?", (contratto_id,)).iloc[0].to_dict()
        tabs_op = st.tabs(["🧩 Servizi", "🛠️ Lavori", "📎 Documenti"])
        with tabs_op[0]:
            try:
                servizi = json.loads(ctr.get("servizi_json") or "[]")
            except Exception:
                servizi = []
            st.dataframe(pd.DataFrame(servizi if servizi else [{"Servizio":"Nessun servizio specificato","Descrizione":""}]), use_container_width=True, hide_index=True)
        with tabs_op[1]:
            lavori_cliente(int(ctr["cliente_id"]), contratto_id)
        with tabs_op[2]:
            documenti_cliente(int(ctr["cliente_id"]), contratto_id)
        return

    df = read_df("SELECT c.id, cl.ragione_sociale cliente, c.titolo, c.data_decorrenza, c.data_scadenza, c.importo_totale, c.stato, COALESCE(s.nome || ' ' || COALESCE(s.cognome,''), 'Non assegnato') staff FROM contratti c JOIN clienti cl ON cl.id=c.cliente_id LEFT JOIN staff s ON s.id=c.staff_id ORDER BY c.id DESC")
    if df.empty: st.info("Nessun contratto."); return
    v=df.copy(); v["importo_totale"]=v["importo_totale"].apply(money); st.dataframe(v, use_container_width=True, hide_index=True)
    opts = {f"ID {r['id']} · {r['cliente']} · {r['titolo']}": int(r['id']) for _,r in df.iterrows()}
    contratto_id = opts[st.selectbox("Seleziona contratto", list(opts.keys()))]
    records = contract_records(contratto_id)
    azienda, cliente, ctr, template, servizi = records
    c1,c2,c3,c4,c5 = st.columns(5)
    staff_nome = read_df("SELECT COALESCE(nome || ' ' || COALESCE(cognome,''), 'Non assegnato') staff FROM staff WHERE id=?", (ctr.get("staff_id"),)) if ctr.get("staff_id") else pd.DataFrame()
    staff_label_card = staff_nome.iloc[0]["staff"] if not staff_nome.empty else "Non assegnato"
    with c1: card("👤","Cliente",cliente["ragione_sociale"])
    with c2: card("📄","Stato", ctr["stato"])
    with c3: card("💰","Importo", money(ctr["importo_totale"]))
    with c4: card("📅","Scadenza", ctr["data_scadenza"])
    with c5: card("👤","Staff", staff_label_card)
    tabs = st.tabs(["✏️ Dati", "🧩 Servizi", "💶 Pagamenti", "🛠️ Lavori", "📎 Documenti", "📄 File"])
    with tabs[0]:
        with st.form("edit_contract"):
            cc1,cc2,cc3 = st.columns(3)
            titolo = cc1.text_input("Titolo", ctr["titolo"]); stato = cc2.selectbox("Stato", ["Bozza","Inviato","Firmato","Attivo","Sospeso","Scaduto","Archiviato"], index=["Bozza","Inviato","Firmato","Attivo","Sospeso","Scaduto","Archiviato"].index(ctr["stato"] if ctr["stato"] in ["Bozza","Inviato","Firmato","Attivo","Sospeso","Scaduto","Archiviato"] else "Bozza")); modalita = cc3.selectbox("Pagamento", ["Mensile","Bimestrale","Trimestrale","Semestrale","Annuale"], index=["Mensile","Bimestrale","Trimestrale","Semestrale","Annuale"].index(ctr["modalita_pagamento"]))
            d1,d2,d3=st.columns(3); decor=d1.date_input("Decorrenza", parse_date(ctr["data_decorrenza"], date.today())); durata=d2.number_input("Durata mesi",1,60,int(ctr["durata_mesi"])); importo=d3.number_input("Importo imponibile",0.0,value=float(ctr["importo_totale"]),step=100.0)
            iva=d1.number_input("IVA %",0.0,100.0,value=float(ctr["iva_percentuale"])); foro=d2.text_input("Foro",ctr["foro_competente"]); rigenera=d3.checkbox("Rigenera rate")
            label_staff, staff_opts = staff_select("Responsabile staff", f"edit_staff_{contratto_id}", include_none=True)
            staff_id = staff_opts[label_staff]
            note=st.text_area("Note", ctr.get("note") or "")
            if st.form_submit_button("Salva modifiche contratto"):
                scad=add_months(decor,int(durata)); execute("UPDATE contratti SET titolo=?,stato=?,modalita_pagamento=?,data_decorrenza=?,data_scadenza=?,durata_mesi=?,importo_totale=?,iva_percentuale=?,foro_competente=?,staff_id=?,note=? WHERE id=?", (titolo,stato,modalita,decor.isoformat(),scad.isoformat(),int(durata),float(importo),float(iva),foro,staff_id,note,contratto_id))
                if rigenera: genera_rate(contratto_id, decor, int(durata), float(importo), float(iva), modalita)
                st.success("Contratto aggiornato."); st.rerun()
    with tabs[1]:
        sdf = pd.DataFrame(servizi if servizi else [{"Servizio":"","Descrizione":""}])
        edited = st.data_editor(sdf, num_rows="dynamic", use_container_width=True, hide_index=True, key=f"serv_{contratto_id}")
        if st.button("Salva servizi"):
            records = [r for r in edited.to_dict("records") if str(r.get("Servizio","")).strip()]
            execute("UPDATE contratti SET servizi_json=? WHERE id=?", (json.dumps(records, ensure_ascii=False), contratto_id)); st.success("Servizi salvati."); st.rerun()
    with tabs[2]:
        pagamenti_contratto(contratto_id)
    with tabs[3]:
        lavori_cliente(int(ctr["cliente_id"]), contratto_id)
    with tabs[4]:
        documenti_cliente(int(ctr["cliente_id"]), contratto_id)
    with tabs[5]:
        if st.button("📝 Rigenera contratto Word/PDF"):
            docx,pdf=generate_contract_files(contratto_id); st.success("Contratto rigenerato."); file_download_button(docx,"Scarica Word"); file_download_button(pdf,"Scarica PDF")
        file_download_button(ctr.get("file_docx"), "Scarica Word esistente"); file_download_button(ctr.get("file_pdf"), "Scarica PDF esistente")
        up=st.file_uploader("Carica contratto firmato", type=["pdf","jpg","jpeg","png"])
        if st.button("Salva firmato") and up:
            rel=save_upload(up,DOC_DIR,f"contratto_firmato_{contratto_id}_"); execute("UPDATE contratti SET file_firmato=? WHERE id=?", (rel,contratto_id)); execute("INSERT INTO documenti (cliente_id,contratto_id,tipo,titolo,file_path,note,created_at) VALUES (?,?,?,?,?,?,?)", (ctr["cliente_id"],contratto_id,"Contratto firmato","Contratto firmato",rel,"Caricato manualmente",now_iso())); st.success("Firmato caricato."); st.rerun()


def elimina_movimento_pagamento(movimento_id: int, pagamento_id: int):
    execute("DELETE FROM incassi_rate WHERE id=? AND pagamento_id=?", (movimento_id, pagamento_id))
    update_stato_pagamento(pagamento_id)


def aggiorna_movimento_pagamento(movimento_id: int, pagamento_id: int, tipo: str, importo: float, data_pagamento: date, allegato_file: str | None, note: str):
    execute("""
        UPDATE incassi_rate
        SET tipo_movimento=?, importo_pagato=?, data_pagamento=?, allegato_file=?, note=?
        WHERE id=? AND pagamento_id=?
    """, (tipo, float(importo), data_pagamento.isoformat(), allegato_file, note, movimento_id, pagamento_id))
    update_stato_pagamento(pagamento_id)


def render_movimenti_pagamento(pid: int, cliente_nome: str, key_prefix: str):
    mov=read_df("""
        SELECT i.*, COALESCE(s.nome || ' ' || COALESCE(s.cognome,''), 'Non indicato') registrato_da
        FROM incassi_rate i
        LEFT JOIN staff s ON s.id=i.registrato_da_staff_id
        WHERE i.pagamento_id=?
        ORDER BY i.data_pagamento DESC, i.id DESC
    """, (pid,))
    st.markdown("### Storico acconti / saldi della rata")
    if mov.empty:
        st.info("Nessun acconto o saldo registrato per questa rata.")
        return
    visual=mov[["id","tipo_movimento","data_pagamento","importo_pagato","registrato_da","allegato_file","note"]].copy()
    visual["importo_pagato"]=visual["importo_pagato"].apply(money)
    st.dataframe(visual, use_container_width=True, hide_index=True)

    st.markdown("### Modifica / elimina movimento")
    opts={f"ID {int(r['id'])} · {r['tipo_movimento']} · {r['data_pagamento']} · {money(r['importo_pagato'])}": int(r['id']) for _,r in mov.iterrows()}
    mid=opts[st.selectbox("Movimento", list(opts.keys()), key=f"{key_prefix}_mov_sel")]
    r=mov[mov.id==mid].iloc[0]
    tipo_opts=["Acconto","Saldo","Pagamento parziale","Rettifica"]
    tipo_index=tipo_opts.index(str(r['tipo_movimento'])) if str(r['tipo_movimento']) in tipo_opts else 0
    data0=to_date(r['data_pagamento']) or date.today()
    allegato_attuale=str(r.get('allegato_file') or '')
    with st.form(f"{key_prefix}_edit_mov_{mid}"):
        c1,c2,c3=st.columns(3)
        tipo=c1.selectbox("Tipo movimento", tipo_opts, index=tipo_index)
        imp=c2.number_input("Importo pagato", min_value=0.0, value=float(r['importo_pagato'] or 0), step=10.0)
        dp=c3.date_input("Data pagamento", value=data0)
        st.caption("Allegato attuale: " + (allegato_attuale if allegato_attuale else "nessun allegato"))
        up=st.file_uploader("Sostituisci allegato PDF/foto", type=["pdf","png","jpg","jpeg","webp"], key=f"{key_prefix}_up_mov_{mid}")
        note=st.text_area("Note movimento", value=str(r.get('note') or ''))
        csave, cdel = st.columns(2)
        salva=csave.form_submit_button("💾 Salva modifica movimento")
        elimina=cdel.form_submit_button("🗑️ Elimina movimento")
        if salva:
            allegato=allegato_attuale or None
            if up:
                allegato=save_upload(up, PAY_DIR, f"pagamento_{pid}_")
            aggiorna_movimento_pagamento(mid, pid, tipo, imp, dp, allegato, note)
            st.success("Movimento aggiornato e saldo rata ricalcolato."); st.rerun()
        if elimina:
            elimina_movimento_pagamento(mid, pid)
            st.warning("Movimento eliminato e saldo rata ricalcolato."); st.rerun()




def render_riepilogo_pagamenti(cliente_id: int | None = None, contratto_id: int | None = None):
    """Riepilogo dinamico pagamenti per cliente/contratto selezionato.
    Mostra imponibile, IVA, lordo, incassi, acconti, residui e media mensile.
    """
    if contratto_id:
        c_where = "WHERE c.id=?"; c_params = (contratto_id,)
        p_where = "WHERE c.id=?"; p_params = (contratto_id,)
        i_where = "WHERE c.id=?"; i_params = (contratto_id,)
    elif cliente_id:
        c_where = "WHERE c.cliente_id=?"; c_params = (cliente_id,)
        p_where = "WHERE c.cliente_id=?"; p_params = (cliente_id,)
        i_where = "WHERE c.cliente_id=?"; i_params = (cliente_id,)
    else:
        c_where = ""; c_params = ()
        p_where = ""; p_params = ()
        i_where = ""; i_params = ()

    contratti = read_df(f"""
        SELECT c.id, c.importo_totale, c.iva_percentuale, c.durata_mesi, c.stato
        FROM contratti c
        {c_where}
    """, c_params)
    rate = payments_df(p_where, p_params)
    incassi = read_df(f"""
        SELECT i.*, p.totale AS totale_rata, c.id AS contratto_id
        FROM incassi_rate i
        JOIN pagamenti p ON p.id=i.pagamento_id
        JOIN contratti c ON c.id=p.contratto_id
        {i_where}
    """, i_params)

    if contratti.empty and rate.empty:
        st.info("Nessun dato pagamento da riepilogare.")
        return

    imponibile_contratti = float(contratti["importo_totale"].fillna(0).sum()) if not contratti.empty else 0.0
    iva_contratti = float((contratti["importo_totale"].fillna(0) * contratti["iva_percentuale"].fillna(0) / 100).sum()) if not contratti.empty else 0.0
    lordo_contratti = imponibile_contratti + iva_contratti
    durata_totale_mesi = int(contratti["durata_mesi"].fillna(0).sum()) if not contratti.empty else 0
    incasso_medio_mensile = (lordo_contratti / durata_totale_mesi) if durata_totale_mesi > 0 else 0.0

    totale_rate_lordo = float(rate["totale"].fillna(0).sum()) if not rate.empty else 0.0
    incassato_totale = float(rate["pagato"].fillna(0).sum()) if not rate.empty else 0.0
    residuo_totale = float(rate["residuo"].fillna(0).clip(lower=0).sum()) if not rate.empty else max(lordo_contratti - incassato_totale, 0.0)
    rate_pagate = rate[(rate["pagato"].fillna(0) + 0.01) >= rate["totale"].fillna(0)] if not rate.empty else rate
    rate_acconto = rate[(rate["pagato"].fillna(0) > 0) & ((rate["pagato"].fillna(0) + 0.01) < rate["totale"].fillna(0))] if not rate.empty else rate
    rate_mancanti = rate[rate["residuo"].fillna(0) > 0.01] if not rate.empty else rate
    oggi_s = date.today().isoformat()
    rate_scadute = rate[(rate["data_scadenza"].astype(str) < oggi_s) & (rate["residuo"].fillna(0) > 0.01)] if not rate.empty else rate

    incassi_importo = float(incassi["importo_pagato"].fillna(0).sum()) if not incassi.empty else 0.0
    acconti_importo = float(incassi[incassi["tipo_movimento"].astype(str).str.lower().str.contains("acconto|parziale", na=False)]["importo_pagato"].fillna(0).sum()) if not incassi.empty else 0.0
    saldi_importo = float(incassi[incassi["tipo_movimento"].astype(str).str.lower().str.contains("saldo", na=False)]["importo_pagato"].fillna(0).sum()) if not incassi.empty else 0.0

    st.markdown("### Riepilogo economico pagamenti")
    a,b,c,d = st.columns(4)
    with a: card("📄", "Totale contratti", money(imponibile_contratti))
    with b: card("🧾", "IVA contratti", money(iva_contratti))
    with c: card("💶", "Totale contratti IVA inclusa", money(lordo_contratti))
    with d: card("📆", "Incasso medio mensile", money(incasso_medio_mensile))

    e,f,g,h = st.columns(4)
    with e: card("✅", "Totale incassato", money(incassato_totale))
    with f: card("🟠", "Totale acconti", money(acconti_importo))
    with g: card("🟢", "Totale saldi", money(saldi_importo))
    with h: card("📉", "Totale rate mancanti", money(residuo_totale))

    i,j,k,l = st.columns(4)
    with i: card("🔢", "Numero rate totali", str(len(rate)) if not rate.empty else "0")
    with j: card("✅", "Rate pagate", f"{len(rate_pagate)} / {money(float(rate_pagate['totale'].sum()) if not rate_pagate.empty else 0)}")
    with k: card("🟠", "Rate con acconto", f"{len(rate_acconto)} / {money(float(rate_acconto['pagato'].sum()) if not rate_acconto.empty else 0)}")
    with l: card("🔴", "Rate scadute", f"{len(rate_scadute)} / {money(float(rate_scadute['residuo'].sum()) if not rate_scadute.empty else 0)}")

    with st.expander("Dettaglio riepilogo", expanded=False):
        dettaglio = pd.DataFrame([
            {"Voce":"Totale contratti imponibile", "Importo": money(imponibile_contratti)},
            {"Voce":"IVA contratti", "Importo": money(iva_contratti)},
            {"Voce":"Totale contratti IVA inclusa", "Importo": money(lordo_contratti)},
            {"Voce":"Totale rate generate", "Importo": money(totale_rate_lordo)},
            {"Voce":"Totale pagamenti/incassi ricevuti", "Importo": money(incassato_totale)},
            {"Voce":"Totale acconti/parziali", "Importo": money(acconti_importo)},
            {"Voce":"Totale saldi", "Importo": money(saldi_importo)},
            {"Voce":"Totale rate mancanti/residuo", "Importo": money(residuo_totale)},
            {"Voce":"Incasso medio mensile previsto", "Importo": money(incasso_medio_mensile)},
        ])
        st.dataframe(dettaglio, use_container_width=True, hide_index=True)
def pagamenti_contratto(contratto_id: int | None = None, cliente_id: int | None = None):
    if contratto_id:
        where = "WHERE p.contratto_id=?"; params=(contratto_id,)
        key = f"contratto_{contratto_id}"
    elif cliente_id:
        where = "WHERE c.cliente_id=?"; params=(cliente_id,)
        key = f"cliente_{cliente_id}"
    else:
        where = ""; params=()
        key = "all"

    df = payments_df(where, params)
    if df.empty:
        st.info("Nessuna rata.")
        return

    st.markdown("### Piano rate")
    f1, f2, f3, f4 = st.columns([1, 1, 1, 2])
    stato_filtro = f1.selectbox("Stato rata", ["Tutte", "Da pagare", "Acconto", "Pagata", "Scaduta", "Sollecitata", "Annullata"], key=f"filtro_stato_{key}")

    anni_disponibili = sorted({str(x)[:4] for x in df["data_scadenza"].dropna().astype(str).tolist() if len(str(x)) >= 4})
    anno_filtro = f2.selectbox("Anno", ["Tutti"] + anni_disponibili, key=f"filtro_anno_{key}")

    mesi_labels = {
        "Tutti": None,
        "01 - Gennaio": "01",
        "02 - Febbraio": "02",
        "03 - Marzo": "03",
        "04 - Aprile": "04",
        "05 - Maggio": "05",
        "06 - Giugno": "06",
        "07 - Luglio": "07",
        "08 - Agosto": "08",
        "09 - Settembre": "09",
        "10 - Ottobre": "10",
        "11 - Novembre": "11",
        "12 - Dicembre": "12",
    }
    mese_label = f3.selectbox("Mese", list(mesi_labels.keys()), key=f"filtro_mese_{key}")
    f4.info("Filtra per anno e mese. La tabella mostra anche il cliente associato quando lavori su più clienti.")

    view = df.copy()
    oggi = date.today().isoformat()
    scad = view["data_scadenza"].astype(str)

    if stato_filtro == "Scaduta":
        view = view[(view["data_scadenza"].astype(str) < oggi) & (view["residuo"].fillna(0) > 0.01)]
    elif stato_filtro != "Tutte":
        view = view[view["stato"].astype(str) == stato_filtro]

    if anno_filtro != "Tutti":
        view = view[view["data_scadenza"].astype(str).str.startswith(str(anno_filtro))]

    mese_val = mesi_labels[mese_label]
    if mese_val:
        view = view[view["data_scadenza"].astype(str).str[5:7] == mese_val]

    tab_s, tab_m, tab_a = st.tabs(["Vista semplice", "Movimenti / allegati", "Modifica avanzata"])

    with tab_s:
        pid = render_payment_simple_table(view if not view.empty else df, key)
        if pid:
            row = df[df.id == pid].iloc[0]
            st.markdown("### Dettaglio rata selezionata")
            r1, r2, r3, r4, r5 = st.columns(5)
            r1.metric("Rata", int(row["numero_rata"]))
            r2.metric("Scadenza", str(row["data_scadenza"]))
            r3.metric("Totale", money(row["totale"]))
            r4.metric("Pagato", money(row["pagato"]))
            r5.metric("Residuo", money(row["residuo"]))
            st.markdown(status_badge_html(row.get("stato"), row.get("residuo"), row.get("data_scadenza")), unsafe_allow_html=True)

    with tab_m:
        st.markdown("### Registra acconto / saldo")
        opts = {f"Rata {int(r['numero_rata'])} · {r['data_scadenza']} · {r['contratto']} · residuo {money(r['residuo'])} · ID {int(r['id'])}": int(r['id']) for _,r in df.iterrows()}
        pid = opts[st.selectbox("Seleziona rata", list(opts.keys()), key=f"pid_{key}")]
        row = df[df.id==pid].iloc[0]
        r1,r2,r3,r4=st.columns(4)
        r1.metric("Totale rata", money(row['totale']))
        r2.metric("Pagato", money(row['pagato']))
        r3.metric("Saldo residuo", money(row['residuo']))
        r4.markdown(status_badge_html(row.get("stato"), row.get("residuo"), row.get("data_scadenza")), unsafe_allow_html=True)

        with st.form(f"incasso_{pid}_{key}"):
            c1,c2,c3=st.columns(3)
            tipo=c1.selectbox("Tipo movimento",["Acconto","Saldo","Pagamento parziale","Rettifica"])
            imp=c2.number_input("Importo pagato",0.0,value=max(float(row["residuo"] or 0),0.0),step=10.0)
            data_pag=c3.date_input("Data pagamento",date.today())
            note=st.text_area("Note incasso")
            up=st.file_uploader("Allega prova pagamento PDF/foto", type=["pdf","png","jpg","jpeg","webp"], key=f"up_incasso_{pid}_{key}")
            if st.form_submit_button("➕ Registra pagamento sulla rata"):
                if imp <= 0:
                    st.error("Inserisci un importo maggiore di zero.")
                else:
                    rel=save_upload(up,PAY_DIR,f"pagamento_{pid}_") if up else None
                    execute("INSERT INTO incassi_rate (pagamento_id,importo_pagato,tipo_movimento,data_pagamento,allegato_file,note,registrato_da_staff_id,created_at) VALUES (?,?,?,?,?,?,?,?)", (pid,float(imp),tipo,data_pag.isoformat(),rel,note,st.session_state.get("staff_user_id"),now_iso()))
                    update_stato_pagamento(pid)
                    st.success("Pagamento registrato e saldo rata aggiornato.")
                    st.rerun()

        render_movimenti_pagamento(pid, str(row.get('cliente') or 'cliente'), f"mov_{key}_{pid}")

    with tab_a:
        st.markdown("### Modifica avanzata rate")
        st.caption("Qui puoi modificare le celle tecniche: rata, scadenza, imponibile, IVA, totale, stato e note.")
        edit = df[["id","contratto","numero_rata","data_scadenza","imponibile","iva","totale","pagato","residuo","stato","note"]].copy()
        edited = st.data_editor(
            edit,
            disabled=["id","contratto","pagato","residuo"],
            num_rows="fixed",
            use_container_width=True,
            hide_index=True,
            key=f"payedit_{key}",
            column_config={
                "stato": st.column_config.SelectboxColumn("stato", options=["Da pagare","Acconto","Pagata","Sollecitata","Annullata"]),
                "data_scadenza": st.column_config.TextColumn("data_scadenza"),
                "imponibile": st.column_config.NumberColumn("imponibile", format="%.2f"),
                "iva": st.column_config.NumberColumn("iva", format="%.2f"),
                "totale": st.column_config.NumberColumn("totale", format="%.2f"),
            }
        )
        if st.button("💾 Salva celle pagamenti", key=f"savepay_{key}"):
            for _,r in edited.iterrows():
                execute(
                    "UPDATE pagamenti SET numero_rata=?,data_scadenza=?,imponibile=?,iva=?,totale=?,stato=?,note=? WHERE id=?",
                    (int(r["numero_rata"]),str(r["data_scadenza"]),float(r["imponibile"]),float(r["iva"]),float(r["totale"]),str(r["stato"]),str(r.get("note") or ""),int(r["id"]))
                )
                update_stato_pagamento(int(r["id"]))
            st.success("Rate aggiornate.")
            st.rerun()


def page_pagamenti():
    header(); section("💶","Pagamenti","Gestione per cliente/contratto con celle editabili, acconti, saldi, movimenti e allegati.")
    df=read_df("SELECT DISTINCT cl.id, cl.ragione_sociale FROM clienti cl JOIN contratti c ON c.cliente_id=cl.id ORDER BY cl.ragione_sociale")
    if df.empty: st.info("Nessun cliente con contratti."); return
    opts={"Tutti i clienti": None} | {f"{r['ragione_sociale']} · ID {r['id']}": int(r['id']) for _,r in df.iterrows()}
    cid=opts[st.selectbox("Cliente", list(opts.keys()))]

    if cid is None:
        ctrs=read_df("SELECT id,titolo FROM contratti ORDER BY id DESC")
    else:
        ctrs=read_df("SELECT id,titolo FROM contratti WHERE cliente_id=? ORDER BY id DESC", (cid,))

    opts2={"Tutti i contratti":None} | {f"ID {r['id']} · {r['titolo']}": int(r['id']) for _,r in ctrs.iterrows()}
    ct=opts2[st.selectbox("Contratto", list(opts2.keys()))]
    render_riepilogo_pagamenti(cliente_id=cid, contratto_id=ct)
    st.divider()
    pagamenti_contratto(contratto_id=ct, cliente_id=None if ct else cid)

def page_lavori():
    header(); section("🛠️","Lavori","Storico operativo collegabile a cliente, contratto e allegati.")
    label,opts=client_select("Cliente", "lav_cli")
    if not opts: st.info("Nessun cliente."); return
    lavori_cliente(opts[label])

def page_documenti():
    header(); section("📎","Archivio documenti","Documenti allegati a clienti e contratti.")
    label,opts=client_select("Cliente", "doc_cli")
    if not opts: st.info("Nessun cliente."); return
    documenti_cliente(opts[label])

def next_invoice_number() -> tuple[str,int]:
    year=date.today().year
    df=read_df("SELECT COUNT(*) n FROM fatture WHERE anno=?", (year,))
    return f"{int(df.iloc[0]['n'])+1}/{year}", year

def page_fatture():
    header(); section("🧾","Fatture interne V1","PDF di cortesia, numerazione e collegamento a clienti/rate/lavori. No SdI/XML per ora.")
    tab1,tab2=st.tabs(["Nuova fattura", "Archivio"])
    with tab1:
        label, azi_opts=company_select("Azienda emittente", "fat_az"); azienda_id=azi_opts[label]
        cl_label, cl_opts=client_select("Cliente", "fat_cl")
        if not cl_opts: st.info("Inserisci prima un cliente."); return
        cliente_id=cl_opts[cl_label]
        num,year=next_invoice_number()

        # Campi fuori dal form: così IVA e totale si aggiornano in tempo reale mentre modifichi imponibile/IVA.
        c1,c2,c3=st.columns(3)
        numero=c1.text_input("Numero", num, key="fat_numero")
        data_f=c2.date_input("Data fattura",date.today(), key="fat_data")
        scad=c3.date_input("Scadenza", date.today()+timedelta(days=30), key="fat_scad")
        desc=st.text_area("Descrizione", "Canone/servizi professionali come da contratto.", key="fat_desc")

        c_imp, c_iva = st.columns(2)
        imponibile=c_imp.number_input("Imponibile",0.0,value=0.0,step=100.0,key="fat_imponibile")
        iva_pct=c_iva.number_input("IVA %",0.0,100.0,value=22.0,key="fat_iva_pct")
        iva=round(float(imponibile)*float(iva_pct)/100,2)
        totale=round(float(imponibile)+iva,2)

        m1,m2,m3=st.columns(3)
        m1.metric("Imponibile", money(float(imponibile)))
        m2.metric("IVA", money(iva))
        m3.metric("Totale fattura", money(totale))

        stato=st.selectbox("Stato", ["Bozza","Pronta","Inviata","Pagata","Scartata","Stornata"], key="fat_stato")
        note=st.text_area("Note", key="fat_note")

        if st.button("Crea fattura interna e PDF", key="btn_crea_fattura"):
            if float(imponibile) <= 0:
                st.error("Inserisci un imponibile maggiore di zero.")
            else:
                fid=execute("INSERT INTO fatture (azienda_id,cliente_id,numero,anno,data_fattura,scadenza,descrizione,imponibile,iva_percentuale,iva,totale,stato,note,emessa_da_staff_id,created_at) VALUES (?,?,?,?,?,?,?,?,?,?,?,?,?,?,?)", (azienda_id,cliente_id,numero,year,data_f.isoformat(),scad.isoformat(),desc,float(imponibile),float(iva_pct),iva,totale,stato,note,st.session_state.get("staff_user_id"),now_iso()))
                rel=generate_invoice_pdf(fid)
                st.session_state["ultima_fattura_pdf"] = rel
                st.success("Fattura creata.")

        if st.session_state.get("ultima_fattura_pdf"):
            file_download_button(st.session_state["ultima_fattura_pdf"], "Scarica PDF fattura")
    with tab2:
        df=read_df("""
            SELECT f.id,f.numero,f.data_fattura,cl.ragione_sociale cliente,
                   COALESCE(s.nome || ' ' || COALESCE(s.cognome,''), 'Non indicato') emessa_da,
                   f.descrizione,f.totale,f.stato,f.file_pdf
            FROM fatture f
            JOIN clienti cl ON cl.id=f.cliente_id
            LEFT JOIN staff s ON s.id=f.emessa_da_staff_id
            ORDER BY f.id DESC
        """)
        if df.empty:
            st.info("Nessuna fattura.")
        else:
            visual=df.copy()
            visual["totale"]=visual["totale"].apply(money)
            st.dataframe(visual,use_container_width=True,hide_index=True)

            st.divider()
            section("✏️", "Modifica / elimina fattura", "Seleziona una fattura dall'archivio, modifica i dati, rigenera il PDF oppure elimina la bozza.")

            opts={f"ID {int(r['id'])} · {r['numero']} · {r['cliente']} · {money(r['totale'])}": int(r['id']) for _,r in df.iterrows()}
            chosen=st.selectbox("Fattura da modificare", list(opts.keys()), key="edit_fattura_select")
            fid=opts[chosen]
            fatt=read_df("SELECT * FROM fatture WHERE id=?", (fid,))
            if not fatt.empty:
                f=fatt.iloc[0]
                emessa_da = read_df("SELECT COALESCE(nome || ' ' || COALESCE(cognome,''), username, 'Non indicato') nome_staff FROM staff WHERE id=?", (f.get("emessa_da_staff_id"),))
                st.caption("Fattura emessa da: " + (emessa_da.iloc[0]["nome_staff"] if not emessa_da.empty else "Non indicato"))

                aziende=read_df("SELECT id,nome,piva FROM aziende ORDER BY nome")
                az_opts={f"{r['nome']} · P.IVA {r['piva']} · ID {int(r['id'])}": int(r['id']) for _,r in aziende.iterrows()} if not aziende.empty else {}
                clienti=read_df("SELECT id,ragione_sociale,partita_iva FROM clienti ORDER BY ragione_sociale")
                cl_opts={f"{r['ragione_sociale']} · P.IVA {r.get('partita_iva') or '-'} · ID {int(r['id'])}": int(r['id']) for _,r in clienti.iterrows()} if not clienti.empty else {}

                c1,c2=st.columns(2)
                if az_opts:
                    az_labels=list(az_opts.keys())
                    az_current=next((k for k,v in az_opts.items() if v==int(f.get('azienda_id') or 0)), az_labels[0])
                    azienda_label=c1.selectbox("Azienda emittente", az_labels, index=az_labels.index(az_current), key=f"edit_fat_azienda_{fid}")
                    azienda_id=az_opts[azienda_label]
                else:
                    azienda_id=int(f.get('azienda_id') or 0)
                if cl_opts:
                    cl_labels=list(cl_opts.keys())
                    cl_current=next((k for k,v in cl_opts.items() if v==int(f.get('cliente_id') or 0)), cl_labels[0])
                    cliente_label=c2.selectbox("Cliente", cl_labels, index=cl_labels.index(cl_current), key=f"edit_fat_cliente_{fid}")
                    cliente_id=cl_opts[cliente_label]
                else:
                    cliente_id=int(f.get('cliente_id') or 0)

                c1,c2,c3=st.columns(3)
                numero=c1.text_input("Numero", str(f.get('numero') or ""), key=f"edit_fat_numero_{fid}")
                data_f=c2.date_input("Data fattura", to_date(f.get('data_fattura'), date.today()) or date.today(), key=f"edit_fat_data_{fid}")
                scad=c3.date_input("Scadenza", to_date(f.get('scadenza'), date.today()+timedelta(days=30)) or date.today()+timedelta(days=30), key=f"edit_fat_scad_{fid}")
                desc=st.text_area("Descrizione", str(f.get('descrizione') or ""), key=f"edit_fat_desc_{fid}")

                c_imp,c_iva=st.columns(2)
                imponibile=c_imp.number_input("Imponibile", min_value=0.0, value=float(f.get('imponibile') or 0), step=100.0, key=f"edit_fat_imponibile_{fid}")
                iva_pct=c_iva.number_input("IVA %", min_value=0.0, max_value=100.0, value=float(f.get('iva_percentuale') or 22), step=1.0, key=f"edit_fat_iva_{fid}")
                iva=round(float(imponibile)*float(iva_pct)/100,2)
                totale=round(float(imponibile)+iva,2)
                m1,m2,m3=st.columns(3)
                m1.metric("Imponibile", money(imponibile))
                m2.metric("IVA", money(iva))
                m3.metric("Totale", money(totale))

                stato_opts=["Bozza","Pronta","Inviata","Pagata","Scartata","Stornata"]
                stato_val=str(f.get('stato') or "Bozza")
                stato_idx=stato_opts.index(stato_val) if stato_val in stato_opts else 0
                stato=st.selectbox("Stato", stato_opts, index=stato_idx, key=f"edit_fat_stato_{fid}")
                note=st.text_area("Note", str(f.get('note') or ""), key=f"edit_fat_note_{fid}")

                b1,b2,b3=st.columns(3)
                if b1.button("💾 Salva modifiche", key=f"save_fattura_{fid}"):
                    if not numero.strip():
                        st.error("Il numero fattura è obbligatorio.")
                    elif float(imponibile) <= 0:
                        st.error("Inserisci un imponibile maggiore di zero.")
                    else:
                        execute("UPDATE fatture SET azienda_id=?, cliente_id=?, numero=?, anno=?, data_fattura=?, scadenza=?, descrizione=?, imponibile=?, iva_percentuale=?, iva=?, totale=?, stato=?, note=? WHERE id=?", (azienda_id,cliente_id,numero,int(data_f.year),data_f.isoformat(),scad.isoformat(),desc,float(imponibile),float(iva_pct),iva,totale,stato,note,fid))
                        st.success("Fattura aggiornata.")
                        st.rerun()

                if b2.button("📄 Rigenera PDF", key=f"regen_fattura_{fid}"):
                    execute("UPDATE fatture SET azienda_id=?, cliente_id=?, numero=?, anno=?, data_fattura=?, scadenza=?, descrizione=?, imponibile=?, iva_percentuale=?, iva=?, totale=?, stato=?, note=? WHERE id=?", (azienda_id,cliente_id,numero,int(data_f.year),data_f.isoformat(),scad.isoformat(),desc,float(imponibile),float(iva_pct),iva,totale,stato,note,fid))
                    rel=generate_invoice_pdf(fid)
                    st.session_state[f"pdf_fattura_archivio_{fid}"]=rel
                    st.success("PDF rigenerato.")

                if b3.button("🗑️ Elimina fattura", key=f"delete_fattura_{fid}"):
                    execute("DELETE FROM fatture WHERE id=?", (fid,))
                    st.warning("Fattura eliminata.")
                    st.rerun()

                pdf_rel=st.session_state.get(f"pdf_fattura_archivio_{fid}") or str(f.get('file_pdf') or "")
                if pdf_rel:
                    file_download_button(pdf_rel, "⬇️ Scarica PDF fattura selezionata")



def estrai_testo_template_da_upload(uploaded: Any) -> str:
    """Estrae testo da PDF/DOCX/TXT per creare template contrattuali modificabili."""
    if not uploaded:
        return ""
    name = (getattr(uploaded, "name", "") or "").lower()
    data = get_upload_bytes(uploaded)
    if not data:
        return ""
    try:
        if name.endswith(".pdf"):
            return estrai_testo_da_upload_contratto(uploaded)
        if name.endswith(".docx"):
            doc = Document(BytesIO(data))
            parts = []
            for par in doc.paragraphs:
                if par.text and par.text.strip():
                    parts.append(par.text.strip())
            for tbl in doc.tables:
                for row in tbl.rows:
                    vals = [cell.text.strip() for cell in row.cells if cell.text and cell.text.strip()]
                    if vals:
                        parts.append(" | ".join(vals))
            return "\n".join(parts)
        if name.endswith(".txt"):
            return data.decode("utf-8", errors="ignore")
    except Exception as exc:
        raise RuntimeError(f"Non sono riuscito a leggere il file template: {exc}") from exc
    return ""


def pulisci_testo_template_contratto(testo: str, ignora_anagrafiche: bool = True, servizi_dinamici: bool = True) -> str:
    """Prepara una base contrattuale riutilizzabile: elimina parti/anagrafiche e lascia il contenuto giuridico operativo."""
    txt = compact_pdf_text_repair(testo or "")
    txt = re.sub(r"\r", "\n", txt)
    txt = re.sub(r"[ \t]+", " ", txt)
    txt = re.sub(r"\n{3,}", "\n\n", txt).strip()

    if ignora_anagrafiche:
        # Se il contratto contiene intestazioni con ragioni sociali/parti, usiamo come base solo dagli articoli in avanti.
        m = re.search(r"(?i)\bART\.?\s*1\b|\bArt\.?\s*1\b", txt)
        if m:
            txt = txt[m.start():].strip()
        # Rimuove blocchi finali di firma/prospetto che il generatore crea già con i dati reali del nuovo cliente.
        stop = re.search(r"(?i)\bPROSPETTO\s+MODALIT[ÀA]\s+DI\s+PAGAMENTO\b|\bLETTO,?\s+CONFERMATO\s+E\s+SOTTOSCRITTO\b|\bFIRME\b", txt)
        if stop:
            txt = txt[:stop.start()].strip()
        # Rimuove righe anagrafiche residue nel caso di testi senza Art. 1.
        righe = []
        for line in txt.splitlines():
            l = line.strip()
            if re.search(r"(?i)(P\.?\s*IVA|C\.?F\.?|codice\s+fiscale|PEC|REA|sede\s+legale|ragione\s+sociale|cliente\s*:|consulente\s*:)", l):
                continue
            righe.append(line)
        txt = "\n".join(righe).strip()

    if servizi_dinamici:
        # Sostituisce il contenuto dell'ART. 2 con il segnaposto dei servizi selezionabili/liberi.
        pat = re.compile(r"(?is)(ART\.?\s*2\s*[–—-].*?)(?=\n?\s*ART\.?\s*3\s*[–—-])")
        if pat.search(txt):
            txt = pat.sub("ART. 2 – SERVIZI COMPRESI\n{{SERVIZI}}\n\n", txt)
        elif "{{SERVIZI}}" not in txt:
            txt += "\n\nART. 2 – SERVIZI COMPRESI\n{{SERVIZI}}"

    txt = re.sub(r"\n{3,}", "\n\n", txt).strip()
    return txt

def page_template():
    header(); section("📄","Template contratti","Base contrattuale modificabile, riutilizzabile e pulita dai dati anagrafici dei clienti.")
    df=read_df("SELECT * FROM templates_contratto ORDER BY id")
    if not df.empty:
        st.dataframe(df[["id","nome","descrizione","attivo"]], use_container_width=True, hide_index=True)
    else:
        st.info("Nessun template presente. Crea il primo template.")

    tab_new, tab_edit = st.tabs(["➕ Crea nuovo template", "✏️ Modifica template esistente"])

    with tab_new:
        st.markdown("### Crea nuovo template")
        st.caption("Puoi copiare e incollare il testo oppure allegare un file Word/PDF. Le ragioni sociali, le P.IVA, PEC e dati delle parti non vengono usati come base: il template conserva solo il contenuto contrattuale.")
        fonte = st.radio("Fonte template", ["Copia/incolla testo", "Allega Word/PDF/TXT"], horizontal=True, key="tpl_new_source")
        uploaded = None
        testo_importato = ""
        if fonte == "Allega Word/PDF/TXT":
            uploaded = st.file_uploader("Carica base contrattuale", type=["pdf","docx","txt"], key="tpl_upload_new")
            col_a, col_b = st.columns([1,2])
            if col_a.button("📖 Leggi file template", key="read_tpl_new"):
                if not uploaded:
                    st.error("Carica prima un file Word, PDF o TXT.")
                else:
                    try:
                        st.session_state["tpl_new_raw_text"] = estrai_testo_template_da_upload(uploaded)
                        st.success("File letto. Controlla e salva il template.")
                    except Exception as exc:
                        st.error(str(exc))
            testo_importato = st.session_state.get("tpl_new_raw_text", "")
        else:
            testo_importato = st.session_state.get("tpl_new_raw_text", "")

        ignora = st.checkbox("Ignora ragioni sociali, P.IVA, PEC e dati delle parti", value=True, key="tpl_new_clean_parties")
        servizi_dyn = st.checkbox("Rendi dinamico l'ART. 2 / Servizi con il segnaposto {{SERVIZI}}", value=True, key="tpl_new_dyn_services")

        raw_text = st.text_area("Testo base da usare", value=testo_importato, height=360, key="tpl_new_text_area")
        if st.button("🧹 Pulisci testo template", key="clean_tpl_new"):
            st.session_state["tpl_new_raw_text"] = pulisci_testo_template_contratto(raw_text, ignora, servizi_dyn)
            st.rerun()

        with st.form("template_new_form"):
            nome=st.text_input("Nome template", value="")
            desc=st.text_input("Descrizione", value="")
            testo=st.text_area("Testo finale template", value=pulisci_testo_template_contratto(raw_text, ignora, servizi_dyn) if raw_text else "", height=520)
            attivo=st.checkbox("Attivo", value=True)
            salva=st.form_submit_button("💾 Crea nuovo template")
            if salva:
                if not nome.strip():
                    st.error("Inserisci il nome del template.")
                elif not testo.strip():
                    st.error("Inserisci o importa il testo del template.")
                else:
                    execute("INSERT INTO templates_contratto (nome,descrizione,testo_base,attivo,created_at) VALUES (?,?,?,?,?)", (nome.strip(),desc.strip(),testo.strip(),1 if attivo else 0,now_iso()))
                    st.session_state.pop("tpl_new_raw_text", None)
                    st.success("Nuovo template creato."); st.rerun()

    with tab_edit:
        if df.empty:
            st.info("Non ci sono template da modificare.")
            return
        ids=df["id"].astype(int).tolist()
        choice=st.selectbox("Template da modificare", ids, format_func=lambda x:f"ID {x} - {df[df.id==x].iloc[0]['nome']}")
        row=df[df.id==choice].iloc[0].to_dict()
        st.markdown("### Modifica template esistente")
        up_edit = st.file_uploader("Sostituisci testo da Word/PDF/TXT", type=["pdf","docx","txt"], key=f"tpl_upload_edit_{choice}")
        col1,col2,col3 = st.columns(3)
        if col1.button("📖 Leggi file e sostituisci testo", key=f"read_tpl_edit_{choice}"):
            if not up_edit:
                st.error("Carica prima un file.")
            else:
                try:
                    st.session_state[f"tpl_edit_text_{choice}"] = pulisci_testo_template_contratto(estrai_testo_template_da_upload(up_edit), True, True)
                    st.success("Testo estratto e pulito. Controlla prima di salvare.")
                except Exception as exc:
                    st.error(str(exc))
        if col2.button("🧹 Pulisci anagrafiche dal testo attuale", key=f"clean_tpl_edit_{choice}"):
            st.session_state[f"tpl_edit_text_{choice}"] = pulisci_testo_template_contratto(st.session_state.get(f"tpl_edit_text_{choice}", row.get("testo_base", BASE_CONTRATTO)), True, True)
            st.rerun()
        if col3.button("↩️ Ripristina testo salvato", key=f"reset_tpl_edit_{choice}"):
            st.session_state.pop(f"tpl_edit_text_{choice}", None); st.rerun()

        edit_text_default = st.session_state.get(f"tpl_edit_text_{choice}", row.get("testo_base", BASE_CONTRATTO))
        with st.form("template_edit_form"):
            nome=st.text_input("Nome template", row.get("nome", ""))
            desc=st.text_input("Descrizione", row.get("descrizione", ""))
            testo=st.text_area("Testo base", edit_text_default, height=560)
            attivo=st.checkbox("Attivo", bool(row.get("attivo",1)))
            csave, cdel = st.columns(2)
            if csave.form_submit_button("💾 Salva modifiche template"):
                execute("UPDATE templates_contratto SET nome=?,descrizione=?,testo_base=?,attivo=? WHERE id=?", (nome.strip(),desc.strip(),testo.strip(),1 if attivo else 0,choice))
                st.session_state.pop(f"tpl_edit_text_{choice}", None)
                st.success("Template aggiornato."); st.rerun()
            if cdel.form_submit_button("🗑️ Disattiva template"):
                execute("UPDATE templates_contratto SET attivo=0 WHERE id=?", (choice,))
                st.warning("Template disattivato."); st.rerun()


def page_admin():
    header(); section("👤","Admin / Staff","Crea membri dello staff e assegna i contratti ai responsabili operativi.")
    tab1, tab2 = st.tabs(["Membri staff", "Assegna contratti"])

    with tab1:
        staff_df = read_df("SELECT * FROM staff ORDER BY stato DESC, nome, cognome")
        if staff_df.empty:
            st.info("Nessun membro staff creato.")
        else:
            st.dataframe(staff_df[["id","username","nome","cognome","ruolo","email","telefono","access_level","is_admin","stato","note"]], use_container_width=True, hide_index=True)

        st.markdown("### Crea / modifica membro staff")
        ids = [0] + staff_df["id"].astype(int).tolist() if not staff_df.empty else [0]
        scelta = st.selectbox("Membro da modificare oppure nuovo", ids, format_func=lambda x: "Nuovo membro staff" if x == 0 else f"ID {x} - {staff_df[staff_df.id==x].iloc[0]['nome']} {staff_df[staff_df.id==x].iloc[0]['cognome'] or ''}")
        row = {} if scelta == 0 else staff_df[staff_df.id == scelta].iloc[0].to_dict()

        with st.form("staff_form"):
            c1,c2,c3 = st.columns(3)
            username = c1.text_input("Username *", row.get("username", "") or "")
            nome = c1.text_input("Nome *", row.get("nome", ""))
            cognome = c1.text_input("Cognome", row.get("cognome", "") or "")
            ruolo = c2.text_input("Ruolo", row.get("ruolo", "") or "")
            email = c2.text_input("Email", row.get("email", "") or "")
            telefono = c3.text_input("Telefono", row.get("telefono", "") or "")
            stato = c3.selectbox("Stato", ["Attivo","Sospeso","Ex staff"], index=["Attivo","Sospeso","Ex staff"].index(row.get("stato","Attivo") if row.get("stato","Attivo") in ["Attivo","Sospeso","Ex staff"] else "Attivo"))
            livelli = ["Operativo Base", "Operativo Avanzato", "Manager Operativo", "Gestione Finanziaria", "Amministratore"]
            livello_attuale = row.get("access_level", "Operativo Base") or "Operativo Base"
            livello_idx = livelli.index(livello_attuale) if livello_attuale in livelli else 0
            access_level = st.selectbox("Livello accesso", livelli, index=livello_idx, help="Operativo Base: solo clienti assegnati, lavori e documenti. Operativo Avanzato: tutti i clienti/lavori/documenti. Manager Operativo: contratti operativi senza pagamenti/fatture. Gestione Finanziaria: clienti, pagamenti, fatture e documenti. Amministratore: vede tutto.")
            is_admin = (access_level == "Amministratore")
            nuova_password = st.text_input("Nuova password / password iniziale", type="password", help="Compila solo se vuoi impostare o cambiare la password.")
            note = st.text_area("Note", row.get("note", "") or "")
            if st.form_submit_button("💾 Salva membro staff"):
                if not username.strip() or not nome.strip():
                    st.error("Username e nome sono obbligatori.")
                elif scelta == 0:
                    if not nuova_password:
                        st.error("Per un nuovo membro staff devi impostare una password iniziale.")
                    else:
                        execute("INSERT INTO staff (username,nome,cognome,ruolo,email,telefono,password_hash,access_level,is_admin,stato,note,created_at) VALUES (?,?,?,?,?,?,?,?,?,?,?,?)", (username,nome,cognome,ruolo,email,telefono,hash_password(nuova_password),access_level,1 if is_admin else 0,stato,note,now_iso()))
                        st.success("Membro staff creato.")
                        st.rerun()
                else:
                    if nuova_password:
                        execute("UPDATE staff SET username=?,nome=?,cognome=?,ruolo=?,email=?,telefono=?,password_hash=?,access_level=?,is_admin=?,stato=?,note=? WHERE id=?", (username,nome,cognome,ruolo,email,telefono,hash_password(nuova_password),access_level,1 if is_admin else 0,stato,note,scelta))
                    else:
                        execute("UPDATE staff SET username=?,nome=?,cognome=?,ruolo=?,email=?,telefono=?,access_level=?,is_admin=?,stato=?,note=? WHERE id=?", (username,nome,cognome,ruolo,email,telefono,access_level,1 if is_admin else 0,stato,note,scelta))
                    st.success("Membro staff aggiornato.")
                    st.rerun()

        if scelta != 0:
            st.warning("Eliminando un membro staff, i contratti assegnati resteranno senza responsabile.")
            if st.button("🗑️ Elimina membro staff", key=f"delete_staff_{scelta}"):
                execute("UPDATE contratti SET staff_id=NULL WHERE staff_id=?", (scelta,))
                execute("DELETE FROM staff WHERE id=?", (scelta,))
                st.success("Membro staff eliminato e contratti scollegati.")
                st.rerun()

    with tab2:
        df = read_df("""
            SELECT c.id, cl.ragione_sociale cliente, c.titolo, c.data_decorrenza, c.data_scadenza,
                   c.importo_totale, c.stato,
                   COALESCE(s.nome || ' ' || COALESCE(s.cognome,''), 'Non assegnato') staff
            FROM contratti c
            JOIN clienti cl ON cl.id=c.cliente_id
            LEFT JOIN staff s ON s.id=c.staff_id
            ORDER BY c.id DESC
        """)
        if df.empty:
            st.info("Nessun contratto disponibile.")
        else:
            v = df.copy()
            v["importo_totale"] = v["importo_totale"].apply(money)
            st.dataframe(v, use_container_width=True, hide_index=True)
            opts = {f"ID {r['id']} · {r['cliente']} · {r['titolo']}": int(r['id']) for _, r in df.iterrows()}
            contratto_id = opts[st.selectbox("Contratto da assegnare", list(opts.keys()), key="admin_assign_contract")]
            label, staff_opts = staff_select("Assegna a membro staff", "admin_assign_staff", include_none=True)
            staff_id = staff_opts[label]
            if st.button("✅ Salva assegnazione contratto"):
                execute("UPDATE contratti SET staff_id=? WHERE id=?", (staff_id, contratto_id))
                st.success("Contratto assegnato.")
                st.rerun()


def page_impostazioni():
    header(); section("⚙️","Impostazioni","Database, cartelle e informazioni tecniche.")
    st.write(f"Database: `{DB_PATH}`")
    st.write(f"Cartella contratti: `{GENERATED_DIR}`")
    st.write(f"Cartella allegati: `{BASE_DIR / 'allegati'}`")
    st.info("Versione Streamlit Pro rifattibile: aziende, clienti, contratti, template, documenti, lavori, pagamenti e fatture sono modificabili.")

# -----------------------------
# Main
# -----------------------------
def main():
    init_db()
    css()
    if not login_screen():
        return
    page=sidebar()
    admin_only = {"Importa contratto","Aziende","Crea nuovo contratto","Template","Impostazioni","Admin"}
    if page in admin_only and not user_is_admin():
        st.error("Questa sezione è riservata all'amministratore.")
        return
    if page in {"Pagamenti","Fatture"} and not can_manage_finance():
        st.error("Questa sezione è riservata alla gestione finanziaria o all'amministratore.")
        return
    if page=="Dashboard": page_dashboard()
    elif page=="Importa contratto": page_importa_contratto()
    elif page=="Aziende": page_aziende()
    elif page=="Clienti CRM": page_clienti()
    elif page=="Crea nuovo contratto": page_crea_contratto()
    elif page=="Contratti": page_contratti()
    elif page=="Pagamenti": page_pagamenti()
    elif page=="Lavori": page_lavori()
    elif page=="Documenti": page_documenti()
    elif page=="Fatture": page_fatture()
    elif page=="Template": page_template()
    elif page=="Impostazioni": page_impostazioni()
    elif page=="Admin": page_admin()

if __name__ == "__main__":
    main()
